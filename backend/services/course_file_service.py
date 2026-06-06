"""
backend/services/course_file_service.py

Generates the complete OBE Course File (.docx) with all 13 sections
using python-docx only — no Node.js dependency.
"""

import io
import json
from pathlib import Path
from typing import Optional

from sqlalchemy import text, select
from sqlalchemy.ext.asyncio import AsyncSession

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.core.logger import get_logger
from backend.core.storage import get_storage
from backend.services.course_service import CourseService

logger = get_logger(__name__)

_CATEGORY = "course_files"

# ── Colours ───────────────────────────────────────────────────────────────────
_NAVY   = (31,  56, 100)
_LIGHT  = (214, 220, 228)
_GREEN  = (226, 239, 218)
_ORANGE = (252, 228, 214)
_WHITE  = (255, 255, 255)
_LGRAY  = (245, 245, 245)
_TEAL   = (0,  112, 130)   # section accent — teal
_AMBER  = (255, 192,  0)   # highlight amber
_STEEL  = (68, 114, 196)   # steel blue accent
_MINT   = (198, 239, 206)  # mint green for CO rows
_PEACH  = (255, 235, 215)  # peach for alternating
_PURPLE = (112,  48, 160)  # purple accent


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: tuple):
    r, g, b = rgb
    hex_color = f"{r:02X}{g:02X}{b:02X}"
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _rgb(tup):
    return RGBColor(*tup)


# ── Style helpers ─────────────────────────────────────────────────────────────

def _run(para, text, bold=False, size=10, color=None, underline=False):
    run = para.add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color or (0, 0, 0))
    run.underline = underline
    return run


def _add_para(doc_or_cell, text="", bold=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=3):
    if hasattr(doc_or_cell, "add_paragraph"):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        _run(p, text, bold=bold, size=size, color=color)
    return p


def _section_title(doc, num, title):
    p = doc.add_paragraph()
    if num > 1:
        run = p.add_run()
        run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)
    # Coloured two-tone banner: teal strip | navy body
    tbl = doc.add_table(rows=1, cols=2)
    tbl.allow_autofit = False
    left = tbl.rows[0].cells[0]
    left.width = Cm(0.35)
    _set_cell_bg(left, _TEAL)
    left.paragraphs[0].paragraph_format.space_before = Pt(5)
    left.paragraphs[0].paragraph_format.space_after  = Pt(5)
    left.paragraphs[0].add_run(" ")
    right = tbl.rows[0].cells[1]
    right.width = Cm(15.65)
    _set_cell_bg(right, _NAVY)
    rp = right.paragraphs[0]
    rp.paragraph_format.space_before = Pt(5)
    rp.paragraph_format.space_after  = Pt(5)
    _run(rp, f"  {num}.  {title}", bold=True, size=13, color=_WHITE)
    # Remove borders using tblPr (compatible with all python-docx versions)
    from docx.oxml import OxmlElement as _OE2
    _tbl_xml = tbl._tbl
    _tbl_pr = _tbl_xml.find(qn("w:tblPr"))
    if _tbl_pr is None:
        _tbl_pr = _OE2("w:tblPr")
        _tbl_xml.insert(0, _tbl_pr)
    _brd = _OE2("w:tblBorders")
    for _side in ("top","left","bottom","right","insideH","insideV"):
        _b = _OE2(f"w:{_side}"); _b.set(qn("w:val"), "none"); _brd.append(_b)
    _tbl_pr.append(_brd)
    doc.add_paragraph()


def _heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _rgb(_TEAL)
    return p


# ── Table builder ─────────────────────────────────────────────────────────────

def _make_table(doc, headers, rows, col_widths_cm):
    num_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=num_cols)
    tbl.style = "Table Grid"

    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Cm(col_widths_cm[i])
        _set_cell_bg(cell, _STEEL)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, hdr, bold=True, size=9, color=_WHITE)

    for ri, row in enumerate(rows):
        bg = _WHITE if ri % 2 == 0 else _LIGHT
        tr = tbl.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            _set_cell_bg(cell, bg)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            _run(p, str(val or ""), size=9)

    return tbl


def _make_header_table(doc, institution_name, institution_address, title, subtitle=""):
    """Creates the header block used in the reference docx (SIT-style merged header)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, _NAVY)
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, institution_name or "Symbiosis Institute of Technology", bold=True, size=11, color=_WHITE)
    if institution_address:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p2, institution_address, size=9, color=_WHITE)
    if title:
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p3, title, bold=True, size=10, color=_WHITE)
    if subtitle:
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p4, subtitle, size=9, color=_WHITE)
    return tbl


# ── Key-safe getter (handles camelCase / snake_case) ─────────────────────────

def _g(row, *keys, default=""):
    """Try multiple key variants; return first non-empty value."""
    for k in keys:
        v = row.get(k)
        if v is not None and str(v).strip():
            return v
    return default


# ── Timetable grid renderer ───────────────────────────────────────────────────

def _render_timetable(doc, timetable: dict):
    """Render the timetable as a proper day x time-slot grid.
    Supports both single-faculty dict and list-of-faculty dicts.
    """
    if not timetable:
        _add_para(doc, "[Timetable not yet uploaded. Upload via the dashboard timetable upload.]",
                  color=(136, 136, 136))
        return

    # Support list of multiple faculty timetables OR a single dict
    if isinstance(timetable, list):
        faculty_tts = timetable
    elif isinstance(timetable, dict) and "faculties" in timetable:
        faculty_tts = timetable["faculties"]
    else:
        faculty_tts = [timetable]

    for tt in faculty_tts:
        faculty = tt.get("faculty_name", "")
        dept    = tt.get("department", "Department of AIML")
        ay      = tt.get("academic_year", "")
        slots   = tt.get("time_slots", [
            "8:45 - 9:40", "9:40 - 10:35", "10:40 - 11:35", "11:35 - 12:30",
            "12:30 - 1:25", "1:25- 2:20", "2:25 - 3:20", "3:25 - 4:20"
        ])
        schedule = tt.get("schedule", {})

        if not schedule and not slots:
            continue

        DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]
        num_cols = 1 + len(slots)

        # ── Outer wrapper table: institute name row ──────────────────────────
        tbl = doc.add_table(rows=1, cols=num_cols)
        tbl.style = "Table Grid"

        for label in [
            "Symbiosis Institute of Technology",
            dept,
            f"Individual Timetable AY {ay}" if ay else "Individual Timetable",
            faculty,
        ]:
            row = tbl.add_row()
            merged = row.cells[0]
            for ci in range(1, num_cols):
                merged = merged.merge(row.cells[ci])
            _set_cell_bg(merged, _NAVY)
            _set_cell_margins(merged, top=60, bottom=60)
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, label, bold=True, size=9, color=_WHITE)

        # Remove the auto-created placeholder first row
        tbl._tbl.remove(tbl.rows[0]._tr)

        # Header row: Day/Time | slot1 | slot2 ...
        hdr_row = tbl.add_row()
        hdr_row.cells[0].width = Cm(2.0)
        _set_cell_bg(hdr_row.cells[0], _NAVY)
        _set_cell_margins(hdr_row.cells[0])
        p = hdr_row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "Day/Time", bold=True, size=8, color=_WHITE)

        slot_w = round(14.0 / max(len(slots), 1), 2)
        for si, slot in enumerate(slots):
            c = hdr_row.cells[si + 1]
            c.width = Cm(slot_w)
            _set_cell_bg(c, _NAVY)
            _set_cell_margins(c, left=40, right=40)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, slot, bold=True, size=7, color=_WHITE)

        def _normalise_day(day_data):
            if isinstance(day_data, dict):
                return day_data
            if isinstance(day_data, list):
                mapping = {}
                for item in day_data:
                    t = item.get("time", "")
                    parts = [p for p in [item.get("course",""), item.get("section",""), item.get("room","")] if p]
                    mapping[t] = "\n".join(parts)
                return mapping
            return {}

        for di, day in enumerate(DAYS):
            raw_day = schedule.get(day, {})
            day_map = _normalise_day(raw_day)
            row = tbl.add_row()
            bg = _WHITE if di % 2 == 0 else _LIGHT

            # Day label cell — 3 sub-rows inside: course, section, room
            row.cells[0].width = Cm(2.0)
            _set_cell_bg(row.cells[0], (220, 225, 235))
            _set_cell_margins(row.cells[0])
            p = row.cells[0].paragraphs[0]
            _run(p, day, bold=True, size=9)

            for si, slot in enumerate(slots):
                entry = day_map.get(slot, "")
                c = row.cells[si + 1]
                c.width = Cm(slot_w)
                _set_cell_bg(c, bg)
                _set_cell_margins(c, left=40, right=40, top=40, bottom=40)
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(p, str(entry or ""), size=7)

        # Summary rows — render each summary item as a separate multi-cell row
        # summary can be:
        #   list of strings  → each string merged across all columns
        #   list of lists    → each inner list = one row of cell values
        summary = tt.get("summary", [])
        if summary:
            for s_item in summary:
                sum_tr = tbl.add_row()
                if isinstance(s_item, list):
                    # Pad/truncate to num_cols
                    vals = list(s_item) + [""] * num_cols
                    vals = vals[:num_cols]
                    for ci, val in enumerate(vals):
                        c = sum_tr.cells[ci]
                        _set_cell_bg(c, _LGRAY)
                        _set_cell_margins(c, top=40, bottom=40, left=60, right=60)
                        p = c.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if ci == num_cols - 1:
                            _run(p, str(val or ""), bold=True, size=8)
                        else:
                            _run(p, str(val or ""), size=8)
                else:
                    # Fall back: merge all cells for a plain text summary line
                    merged_sum = sum_tr.cells[0]
                    for ci in range(1, num_cols):
                        merged_sum = merged_sum.merge(sum_tr.cells[ci])
                    _set_cell_bg(merged_sum, _LGRAY)
                    _set_cell_margins(merged_sum, top=40, bottom=40)
                    p = merged_sum.paragraphs[0]
                    _run(p, str(s_item), size=8)

        doc.add_paragraph()


# ── Main document builder ─────────────────────────────────────────────────────

def _build_docx(data: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = section.right_margin  = Cm(2.5)
        section.top_margin    = section.bottom_margin = Cm(2)

    inst_name = data.get("institution_name") or "Symbiosis Institute of Technology"
    inst_addr = data.get("institution_address") or "SIU Pune 412115, Maharashtra, India"
    dept_full = data.get("department") or "Artificial Intelligence and Machine Learning"
    batch     = data.get("batch", "")

    # ── Cover ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"{data.get('course_name','')} ({data.get('course_code','')}) Course File",
         bold=True, size=16, color=_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"A.Y {data.get('academic_year','')} ({data.get('semester','')} Semester)",
         bold=True, size=13)

    if batch:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"Batch {batch}", bold=True, size=13)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"Department of {dept_full}, {inst_name},", bold=True, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"Symbiosis International (Deemed University) {inst_addr}", bold=True, size=12)

    doc.add_paragraph()

    # Table of Contents
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Course File Contents", bold=True, size=13, color=_NAVY)

    toc_entries = [
        ("1",  "Vision & Mission of the Department"),
        ("2",  "Program Outcomes (POs), Program Educational Objectives (PEOs) and Program Specific Outcomes (PSOs)"),
        ("3",  "Syllabus, Personal Timetable"),
        ("4",  "CO Statements, CO-PO-PSO Mapping with justification"),
        ("5",  "CO Attainment of the course from the previous academic year and the action plan"),
        ("6",  "Session Plan\n  Session Plan with CO mapping to each lecture\n  Text Books\n  Reference books\n"
               "  Web-Links for Online Notes/ YouTube/Coursera/MOOC/SWAYAM/NPTEL Videos/Blogs etc.\n"
               "  Names of Magazines, journals\n  List of Research Articles/Classic papers, review papers\n"
               "  Planning of experiential learning/Guest lectures/Video lectures, industry Visit etc\n"
               "  Tutorial questions with CO mapping"),
        ("7",  "Evaluation plan with CO Mapping\n  For each evaluation component:\n"
               "  The evaluation questions with CO mapping\n"
               "  Marksheet of each Evaluation Component (with a sample copy as a proof)\n"
               "  Final Marksheet"),
        ("8",  "List of Slow and Advanced learners and the action plans"),
        ("9",  "CO Attainment of internal evaluation"),
        ("10", "The reports of the activities planned and conducted"),
        ("11", "Learning Material."),
        ("12", "Question Bank"),
        ("13", "Compiled Attendance"),
    ]
    _make_table(doc,
                ["Sr. No", "Title"],
                [[sr, title] for sr, title in toc_entries],
                [1.5, 14.5])

    # ── 1. Vision & Mission ────────────────────────────────────────────────────
    _section_title(doc, 1, "Vision & Mission of the Department")
    _heading2(doc, "VISION OF THE DEPARTMENT")
    if data.get("vision_text"):
        _add_para(doc, data["vision_text"])
    else:
        _add_para(doc, "To impart quality education with research insights for developing competent global engineers in the field of Artificial Intelligence and Machine Learning to solve societal problems.")

    _heading2(doc, "MISSION OF THE DEPARTMENT")
    if data.get("mission_text"):
        for line in (data["mission_text"] or "").split("\n"):
            if not line.strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            # Detect M1:/M2:/M3: prefix and render it bold
            import re as _re
            m = _re.match(r'^(M\d+:\s*)(.*)', line.strip(), _re.DOTALL)
            if m:
                prefix, rest = m.group(1), m.group(2)
                _run(p, prefix, bold=True, size=10)
                _run(p, rest, size=10)
            else:
                _run(p, line.strip(), size=10)
    else:
        _STANDARD_MISSION = [
            ("M1: ", "To educate students on cutting-edge AIML technologies with strong industry connections to develop problem-solving capabilities, leadership, and teamwork skills."),
            ("M2: ", "To produce quality research through national and international collaborations leading to publications, IPR, and sponsored/funded projects."),
            ("M3: ", "To inculcate professional values with lifelong learning through curricular and co-curricular activities and create globally-aware citizens."),
        ]
        for prefix, rest in _STANDARD_MISSION:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            _run(p, prefix, bold=True, size=10)
            _run(p, rest, size=10)

    # ── 2. POs, PEOs, PSOs ────────────────────────────────────────────────────
    _section_title(doc, 2, "Program Outcomes (POs), Program Educational Objectives (PEOs) and Program Specific Outcomes (PSOs)")
    _heading2(doc, "Program Outcomes (POs)")
    pos = data.get("pos") or []
    _STANDARD_POS = [
        ("PO 1",  "Engineering Knowledge: Apply the knowledge of mathematics, science, engineering fundamentals, and an engineering specialization to the solution of complex engineering problems."),
        ("PO 2",  "Problem analysis: Identify, formulate, review research literature, and analyze complex engineering problems reaching substantiated conclusions using first principles of mathematics, natural sciences, and engineering sciences."),
        ("PO 3",  "Design/development of solutions: Design solutions for complex engineering problems and design system components or processes that meet the specified needs with appropriate consideration for the public health and safety, and the cultural, societal, and environmental considerations."),
        ("PO 4",  "Conduct investigations of complex problems: Use research-based knowledge and research methods including design of experiments, analysis and interpretation of data, and synthesis of the information to provide valid conclusions."),
        ("PO 5",  "Modern tool usage: Create, select, and apply appropriate techniques, resources, and modern engineering and IT tools including prediction and modeling to complex engineering activities with an understanding of the limitations."),
        ("PO 6",  "The engineer and society: Apply reasoning informed by the contextual knowledge to assess societal, health, safety, legal and cultural issues and the consequent responsibilities relevant to the professional engineering practice."),
        ("PO 7",  "Environment and sustainability: Understand the impact of the professional engineering solutions in societal and environmental contexts, and demonstrate the knowledge of, and need for sustainable development."),
        ("PO 8",  "Ethics: Apply ethical principles and commit to professional ethics and responsibilities and norms of the engineering practice."),
        ("PO 9",  "Individual and team work: Function effectively as an individual, and as a member or leader in diverse teams, and in multidisciplinary settings."),
        ("PO 10", "Communication: Communicate effectively on complex engineering activities with the engineering community and with society at large, such as, being able to comprehend and write effective reports and design documentation, make effective presentations, and give and receive clear instructions."),
        ("PO 11", "Project management and finance: Demonstrate knowledge and understanding of the engineering and management principles and apply these to one's own work, as a member and leader in a team, to manage projects and in multidisciplinary environments."),
        ("PO 12", "Life-long learning: Recognize the need for, and have the preparation and ability to engage in independent and life-long learning in the broadest context of technological change."),
    ]
    db_pos_have_text = any(p.get("statement", p.get("description", "")).strip() for p in pos)
    po_rows = (
        [[p.get("po_id",""), p.get("statement", p.get("description",""))] for p in pos]
        if db_pos_have_text
        else [[pid, stmt] for pid, stmt in _STANDARD_POS]
    )
    # Build PO table manually so each row's description can have a bold label prefix
    _po_tbl = doc.add_table(rows=1 + len(po_rows), cols=2)
    _po_tbl.style = "Table Grid"
    for ci, hdr_txt in enumerate(["", "Program Outcomes"]):
        cell = _po_tbl.rows[0].cells[ci]
        cell.width = Cm([1.5, 14.5][ci])
        _set_cell_bg(cell, _NAVY)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, hdr_txt, bold=True, size=9, color=_WHITE)
    import re as _re2
    for ri, (po_id, po_text) in enumerate(po_rows):
        bg = _WHITE if ri % 2 == 0 else _LIGHT
        tr = _po_tbl.rows[ri + 1]
        # ID cell
        id_cell = tr.cells[0]
        id_cell.width = Cm(1.5)
        _set_cell_bg(id_cell, bg)
        _set_cell_margins(id_cell)
        _run(id_cell.paragraphs[0], str(po_id or ""), size=9)
        # Description cell — bold label up to first colon
        desc_cell = tr.cells[1]
        desc_cell.width = Cm(14.5)
        _set_cell_bg(desc_cell, bg)
        _set_cell_margins(desc_cell)
        p = desc_cell.paragraphs[0]
        m = _re2.match(r'^([^:]+:\s*)(.*)', str(po_text or ""), _re2.DOTALL)
        if m:
            _run(p, m.group(1), bold=True, size=9)
            _run(p, m.group(2), size=9)
        else:
            _run(p, str(po_text or ""), size=9)

    _heading2(doc, "Program Educational Objectives (PEOs)")
    peos = [
        ("PEO1", "Apply the knowledge of the latest trends of AIML and will be engaged in technology development and deployment for engineering systems in their profession."),
        ("PEO2", "To be competent AIML engineers with innovative thinking and research attitude to solve the real-world problems."),
        ("PEO3", "To have enhanced interpersonal and managerial skills to function effectively in their profession with social awareness and responsibility."),
    ]
    for pid, ptext in peos:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(f"{pid}: ")
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(ptext)
        r2.font.size = Pt(10)

    _heading2(doc, "Program-specific outcomes (PSOs)")
    psos = [
        ("PSO1", "To apply the concepts of Artificial Intelligence and Machine Learning with practical knowledge in analysis, design and development of intelligent systems and applications to multi-disciplinary problems."),
        ("PSO2", "To provide a concrete foundation to the students in the cutting-edge areas Artificial Intelligence and Machine Learning and excelling in the specialized areas like Natural Language Processing, Computer Vision, Reinforcement Learning, Internet of Things, Cloud computing, Data Security and privacy etc."),
    ]
    _make_table(doc, ["", "Program specific outcomes"], [[pid, ptext] for pid, ptext in psos], [1.5, 14.5])

    # ── 3. Syllabus & Timetable ───────────────────────────────────────────────
    _section_title(doc, 3, "Syllabus, Personal Timetable")

    _heading2(doc, "Individual Timetable:")
    timetable = data.get("timetable") or {}
    _render_timetable(doc, timetable)

    # Additional timetables from attachments
    extra_tt_atts = [a for a in (data.get("attachments") or []) if a.get("section_no") == 3]
    if extra_tt_atts:
        _heading2(doc, "Additional Timetable Uploads")
        for a in extra_tt_atts:
            _add_para(doc, f"Attachment: {a['label']}  ({a['filename']})", size=10)

    # ── List of Students ──────────────────────────────────────────────────────
    # Single unified table, continuous serial numbering, section labels as merged rows,
    # institutional header block at top — matching example exactly.
    _heading2(doc, "List of Students")
    students_all = data.get("students") or []
    if students_all:
        from collections import defaultdict as _dd

        # Institutional header block
        inst_tbl = doc.add_table(rows=1, cols=1)
        inst_tbl.style = "Table Grid"
        inst_hdr_cell = inst_tbl.rows[0].cells[0]
        _set_cell_bg(inst_hdr_cell, _NAVY)
        _set_cell_margins(inst_hdr_cell, top=80, bottom=80)
        p_inst = inst_hdr_cell.paragraphs[0]
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_inst, inst_name, bold=True, size=10, color=_WHITE)
        p_inst2 = inst_hdr_cell.add_paragraph()
        p_inst2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_inst2, f" {batch} Second Year, Sem IV", bold=True, size=9, color=_WHITE)
        p_inst3 = inst_hdr_cell.add_paragraph()
        p_inst3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_inst3, f"Branch - {dept_full} - AIML", bold=True, size=9, color=_WHITE)

        doc.add_paragraph()

        # Group by section, maintain insertion order A → B → C
        by_section = _dd(list)
        for s in students_all:
            sec = (s.get("section") or "").strip().upper() or "All"
            by_section[sec].append(s)

        # One unified table
        unified_tbl = doc.add_table(rows=1, cols=3)
        unified_tbl.style = "Table Grid"
        col_widths = [1.2, 3.5, 11.3]

        hdr = unified_tbl.rows[0]
        for ci, (hdr_txt, w) in enumerate(zip(["SR. No.", "PRN", "AIML"], col_widths)):
            cell = hdr.cells[ci]
            cell.width = Cm(w)
            _set_cell_bg(cell, _NAVY)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, hdr_txt, bold=True, size=9, color=_WHITE)

        global_sr = 1
        for sec_label in sorted(by_section.keys()):
            sec_students = by_section[sec_label]

            # Section label row — merged across all 3 columns
            sec_row = unified_tbl.add_row()
            merged_cell = sec_row.cells[0]
            merged_cell.merge(sec_row.cells[1])
            merged_cell.merge(sec_row.cells[2])
            _set_cell_bg(merged_cell, _LIGHT)
            _set_cell_margins(merged_cell, top=60, bottom=60)
            p_sec = merged_cell.paragraphs[0]
            p_sec.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _run(p_sec, f"Section {sec_label}", bold=True, size=10)

            for s in sec_students:
                data_row = unified_tbl.add_row()
                bg = _WHITE if global_sr % 2 == 1 else _LIGHT
                vals = [str(global_sr), s.get("prn", ""), s.get("name", "")]
                for ci, (val, w) in enumerate(zip(vals, col_widths)):
                    cell = data_row.cells[ci]
                    cell.width = Cm(w)
                    _set_cell_bg(cell, bg)
                    _set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    _run(p, val, size=9)
                global_sr += 1
    else:
        _add_para(doc, "[Student list not available. Add students via the Students page.]",
                  color=(136, 136, 136))

    # ── 4. CO Statements + CO-PO Mapping ─────────────────────────────────────
    _section_title(doc, 4, "CO Statements, CO-PO-PSO Mapping with justification")

    # CO statements as bullet list
    cos = data.get("cos") or []
    if cos:
        _add_para(doc, "COs:", bold=True, size=10)
        for c in cos:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _run(p, c.get("statement", ""), size=10)
        doc.add_paragraph()

    # CO-PO Mapping table with colour-coded strength values
    _heading2(doc, "CO-PO Mapping")
    co_po_matrix = data.get("co_po_matrix") or {}
    po_ids_matrix = [f"PO{i}" for i in range(1, 13)] + ["PSO1", "PSO2"]

    # FIX: use correct NBA standard mapping values (1/2/3) not all-1s.
    # Fall back to standard AIML mapping when DB has no meaningful values.
    _STANDARD_CO_PO = {
        "CO1": {"PO1": 3, "PO2": 3, "PO6": 2, "PSO1": 3, "PSO2": 3},
        "CO2": {"PO1": 3, "PO3": 3, "PO5": 3, "PSO1": 3, "PSO2": 3},
        "CO3": {"PO1": 3, "PO2": 3, "PO3": 3, "PO4": 2, "PSO1": 3, "PSO2": 3},
        "CO4": {"PO1": 3, "PO2": 3, "PO3": 3, "PSO1": 3, "PSO2": 2},
        "CO5": {"PO1": 3, "PO3": 3, "PO5": 3, "PSO1": 3, "PSO2": 3},
    }

    def _matrix_has_real_values(matrix, co_list):
        """Returns True if the DB matrix has values other than 0/1/empty for all POs."""
        for co in co_list:
            mapping = matrix.get(co.get("co_id","")) or {}
            for v in mapping.values():
                try:
                    if int(v) >= 2:
                        return True
                except (ValueError, TypeError):
                    pass
        return False

    use_standard_mapping = not co_po_matrix or not _matrix_has_real_values(co_po_matrix, cos)

    if cos:
        matrix_rows = []
        for co in cos:
            co_id = co.get("co_id", "")
            if use_standard_mapping:
                mapping = _STANDARD_CO_PO.get(co_id, {})
            else:
                mapping = co_po_matrix.get(co_id) or {}

            def _get_val(pid, _mapping=mapping):
                v = (_mapping.get(pid) or
                     _mapping.get(pid.replace("PO", "PO ")) or
                     _mapping.get(pid.replace("PO ", "PO")))
                if v in (None, "", 0, "0"):
                    return ""
                return str(v)

            matrix_rows.append([co_id] + [_get_val(pid) for pid in po_ids_matrix])

        n = len(po_ids_matrix)
        col_w = [1.5] + [round(14.5 / max(n, 1), 2)] * n

        num_cols = 1 + n
        tbl = doc.add_table(rows=1 + len(matrix_rows), cols=num_cols)
        tbl.style = "Table Grid"

        hdr_row = tbl.rows[0]
        for ci, hdr in enumerate(["CO"] + po_ids_matrix):
            cell = hdr_row.cells[ci]
            cell.width = Cm(col_w[ci])
            _set_cell_bg(cell, _NAVY)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, hdr, bold=True, size=8, color=_WHITE)

        strength_bg = {
            "1": (209, 231, 246),
            "2": (130, 188, 235),
            "3": ( 56, 136, 195),
        }
        for ri, row in enumerate(matrix_rows):
            tr = tbl.rows[ri + 1]
            for ci, val in enumerate(row):
                cell = tr.cells[ci]
                cell.width = Cm(col_w[ci])
                _set_cell_margins(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                str_val = str(val) if val not in (None, "") else ""
                if ci == 0:
                    _set_cell_bg(cell, _LGRAY)
                    _run(p, str_val, bold=True, size=8)
                elif str_val in strength_bg:
                    _set_cell_bg(cell, strength_bg[str_val])
                    _run(p, str_val, bold=True, size=8)
                else:
                    _set_cell_bg(cell, _WHITE)
                    _run(p, "", size=8, color=(180, 180, 180))
    else:
        _add_para(doc, "[CO-PO mapping not yet configured. Set it up in Course Setup.]",
                  color=(136, 136, 136))

    doc.add_paragraph()

    # CO-PO Justification — single-column table, one row per paragraph
    co_po_justification = data.get("co_po_justification") or ""

    _STANDARD_JUSTIFICATION = [
        "PO1 (Engineering Knowledge): Strongly linked to all COs because they all require deep understanding of machine learning, algorithms, and data fundamentals.",
        "PO2 (Problem Analysis): Essential for contrasting algorithms, analyzing clustering, and explaining advanced methods.",
        "PO3 (Design/Development of Solutions): Strongly connected to applying and modeling techniques and designing clustering and deep learning solutions.",
        "PO4 (Investigations): Relevant for comparative analysis involving experiments.",
        "PO5 (Modern Tool Usage): Needed for applying dimensionality reduction and deep learning with tools and frameworks.",
        "PO6 (Engineer and Society): Considered moderately for CO1 due to understanding societal impacts related to data and algorithm choices.",
        "PSO1 COs focus on applying core AI/ML concepts such as algorithm understanding, dimensionality reduction, clustering, and deep learning, which directly supports the PSO's emphasis on practical knowledge",
        "CO1 (Contrasting algorithms and data types) lays the foundational understanding required for analyzing and selecting appropriate AI/ML techniques.",
        "CO2 and CO3 (Applying dimensionality reduction and clustering techniques) address designing and developing models\u2014key parts of creating intelligent systems.",
        "CO4 (Explaining advanced clustering for domain-specific datasets) highlights the ability to customize AI solutions for real-world, multi-disciplinary problems.",
        "CO5 (Demonstrating deep learning methods like autoencoders) reflects advanced practical skills needed for developing state-of-the-art intelligent applications.",
        "PSO2- CO1 (Contrast ML algorithm types and classify data types): This CO builds a fundamental understanding of AI/ML concepts by distinguishing various algorithm types and data characteristics, which is essential as a base for advanced, cutting-edge AI areas like NLP and Computer Vision.  CO2 (Apply dimensionality reduction and model unsupervised learning): By applying and modeling core ML techniques, students gain practical skills that are vital for handling high-dimensional data encountered in domains such as IoT and Cloud computing.  CO3 (Model static and hierarchical clustering with comparative analysis): Clustering methods are foundational in many AI applications including Reinforcement Learning and Data Security analytics, thereby supporting specialization in these areas.  CO4 (Explain incremental and advanced clustering algorithms for domain-specific datasets): This CO enables understanding of domain adaptation and specialization, preparing students to tackle specific challenges in emerging AI fields.  CO5 (Demonstrate deep unsupervised learning approaches like autoencoders): Deep learning techniques like autoencoders are crucial in advanced AI areas such as Computer Vision and NLP, directly supporting the PSO\u2019s goal of excelling in cutting-edge domains.",
    ]

    if co_po_justification.strip():
        lines = [l.strip() for l in co_po_justification.split("\n") if l.strip()]
    else:
        lines = _STANDARD_JUSTIFICATION

    _heading2(doc, "Justifications for CO - PO mapping:")
    if lines:
        # Build table with blank spacer rows between each entry (matching example format)
        num_rows = len(lines) * 2 - 1  # content rows + blank spacers
        just_tbl = doc.add_table(rows=num_rows, cols=1)
        just_tbl.style = "Table Grid"
        for ri, line in enumerate(lines):
            # Content row
            cell = just_tbl.rows[ri * 2].cells[0]
            _set_cell_bg(cell, _WHITE)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            _run(p, line, size=9)
            # Blank spacer row (except after last)
            if ri < len(lines) - 1:
                spacer_cell = just_tbl.rows[ri * 2 + 1].cells[0]
                _set_cell_bg(spacer_cell, _WHITE)
                _set_cell_margins(spacer_cell, top=20, bottom=20)
                spacer_cell.paragraphs[0]  # leave empty

    # ── 5. Previous CO Attainment ─────────────────────────────────────────────
    _section_title(doc, 5, "CO Attainment of the course from the previous academic year and the action plan")

    prev_co_att = data.get("prev_co_attainment") or ""
    if prev_co_att.strip():
        # Try to parse as structured CO attainment data (JSON list or plain text)
        import json as _json
        try:
            att_data = _json.loads(prev_co_att)
            if isinstance(att_data, list) and att_data:
                # Render as a table: CO | Target | Attained | Status
                att_hdrs = ["CO", "Target (%)", "Attained (%)", "Status"]
                att_col_w = [2.0, 3.0, 3.0, 8.0]
                att_tbl = doc.add_table(rows=1 + len(att_data), cols=4)
                att_tbl.style = "Table Grid"
                for ci, hdr in enumerate(att_hdrs):
                    cell = att_tbl.rows[0].cells[ci]
                    cell.width = Cm(att_col_w[ci])
                    _set_cell_bg(cell, _NAVY)
                    _set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run(p, hdr, bold=True, size=9, color=_WHITE)
                for ri, entry in enumerate(att_data):
                    bg = _WHITE if ri % 2 == 0 else _LIGHT
                    tr = att_tbl.rows[ri + 1]
                    vals = [
                        entry.get("co", f"CO{ri+1}"),
                        str(entry.get("target", "")),
                        str(entry.get("attained", entry.get("percentage", ""))),
                        entry.get("status", ""),
                    ]
                    for ci, val in enumerate(vals):
                        cell = tr.cells[ci]
                        cell.width = Cm(att_col_w[ci])
                        _set_cell_bg(cell, bg)
                        _set_cell_margins(cell)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 3 else WD_ALIGN_PARAGRAPH.LEFT
                        _run(p, str(val or ""), size=9)
            else:
                for line in prev_co_att.split("\n"):
                    if line.strip():
                        _add_para(doc, line, space_before=2, space_after=2)
        except (ValueError, TypeError):
            # Plain text — render as-is
            for line in prev_co_att.split("\n"):
                if line.strip():
                    _add_para(doc, line, space_before=2, space_after=2)
    else:
        # Blank — matching example (two empty lines, no placeholder text)
        doc.add_paragraph()
        doc.add_paragraph()

    # Action Plan — bold text matching example format
    action_plan = data.get("action_plan") or ""
    if action_plan.strip():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        for line in action_plan.split("\n"):
            if line.strip():
                _run(p, line.strip(), bold=True, size=10)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        _run(p, "Action Plan- The attainment is higher than the set targets for all CO\u2019s.", bold=True, size=10)

    # ── 6. Session Plan ───────────────────────────────────────────────────────
    _section_title(doc, 6, "Session Plan with CO mapping to each lecture")

    # Institutional header block — multi-row, matching example exactly
    sp_tbl = doc.add_table(rows=1, cols=1)
    sp_tbl.style = "Table Grid"
    sp_cell = sp_tbl.rows[0].cells[0]
    _set_cell_bg(sp_cell, _NAVY)
    _set_cell_margins(sp_cell, top=100, bottom=100, left=160, right=160)
    p = sp_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Symbiosis Institute of Technology, Pune", bold=True, size=11, color=_WHITE)
    for line in [
        "Session Plan",
        f"Name of the Department \u2013 {dept_full}",
        f"Name of the course- {data.get('course_name','')}",
        f"Credit - {data.get('credits','3')}",
        f"Semester - {data.get('semester','')}    Batch - {batch}",
        f"Name of the faculty- {data.get('faculty_name','')}",
    ]:
        px = sp_cell.add_paragraph()
        px.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(px, line, size=9, color=_WHITE)

    doc.add_paragraph()

    session_rows = data.get("session_rows") or []
    if session_rows:
        table_data = []
        for i, r in enumerate(session_rows):
            lect    = _g(r, "lect", "lect_no", "lectNo", "lecture_no")
            unit    = _g(r, "unit", "unit_no", "unitNo", "unit_number")
            topic   = _g(r, "topic", "points_to_cover", "pointsToCover", "content", "description")
            method  = _g(r, "method", "methodology", "lecture_method")
            faculty = _g(r, "faculty", "faculty_conducting", "facultyConducting",
                         default=data.get("faculty_name", ""))
            ltype   = _g(r, "type", "lecture_exp_eval", "lectureType", default="Lecture")
            co      = _g(r, "co", "co_mapped", "co_id")
            if isinstance(co, list):
                co = ", ".join(str(x) for x in co)
            table_data.append([str(lect or i+1), str(unit or ""), topic, method, faculty, ltype, co])
        _make_table(doc,
                    ["Lect.\nNo", "Unit No.", "Points to cover", "Methodology",
                     "Faculty Conducting", "Lecture/Exp. Learning/\nEvaluation", "CO"],
                    table_data,
                    [1.2, 1.2, 5.5, 2.2, 3.0, 2.0, 1.4])
    else:
        _add_para(doc, "[Session plan not yet generated. Use the Session Plan page first.]",
                  color=(136, 136, 136))

    # ── Textbooks & Reference Books ───────────────────────────────────────────
    materials = data.get("study_materials") or {}
    textbooks  = materials.get("textbooks") or []
    ref_books  = materials.get("reference_books") or materials.get("references") or []
    web_links  = materials.get("web_links") or materials.get("web") or []
    journals   = materials.get("journals") or []
    moocs      = materials.get("moocs") or []
    research_articles = materials.get("research_articles") or []

    if textbooks or ref_books:
        _heading2(doc, "Textbooks & Reference books/ Beyond Gaps")
        all_books = textbooks + ref_books
        _make_table(doc, ["Book", "Author", "Publisher"],
                    [[b.get("title", b.get("book","") if isinstance(b,dict) else str(b)),
                      b.get("author","") if isinstance(b,dict) else "",
                      b.get("publisher","") if isinstance(b,dict) else ""] for b in all_books],
                    [7.0, 4.0, 5.0])

    if web_links:
        _heading2(doc, "Web-Links for Online Notes/ YouTube/NPTEL Videos/Blogs etc")
        _make_table(doc, ["Sr. No.", "Web Link", "Module"],
                    [[str(i+1),
                      w.get("title", w.get("url","") if isinstance(w,dict) else str(w)),
                      w.get("unit", w.get("module","")) if isinstance(w,dict) else ""]
                     for i, w in enumerate(web_links)],
                    [1.0, 10.0, 5.0])

    if journals:
        _heading2(doc, "Names of Magazines, Journals, E-journals")
        _make_table(doc, ["Sr.No.", "Journal"],
                    [[str(i+1),
                      j.get("title","") if isinstance(j,dict) else str(j)]
                     for i, j in enumerate(journals)],
                    [1.0, 15.0])

    if moocs:
        _heading2(doc, "Recommended MOOC Courses like Coursera / NPTEL / MIT-OCW / edX etc")
        _make_table(doc,
                    ["S.No.", "MOOC Course Link", "Course conducted by", "Course Duration", "Certificate (Y / N)"],
                    [[str(i+1),
                      m.get("title", m.get("url","") if isinstance(m,dict) else str(m)),
                      m.get("platform", m.get("conducted_by","")) if isinstance(m,dict) else "",
                      m.get("duration","") if isinstance(m,dict) else "",
                      m.get("certificate","Y") if isinstance(m,dict) else "Y"]
                     for i, m in enumerate(moocs)],
                    [1.0, 6.0, 3.0, 2.5, 2.0])

    if research_articles:
        _heading2(doc, "List of Research Articles")
        _make_table(doc, ["S.No.", "Research Article Title", "Web Link"],
                    [[str(i+1),
                      a.get("title","") if isinstance(a,dict) else str(a),
                      a.get("url","") if isinstance(a,dict) else ""]
                     for i, a in enumerate(research_articles)],
                    [1.0, 9.0, 6.0])

    # Prepared By / Approved By footer
    if data.get("faculty_name"):
        doc.add_paragraph()
        _add_para(doc, f"Prepared By: {data.get('faculty_name','')}", size=10)
    approved_by = data.get("hod_name") or ""
    if approved_by:
        _add_para(doc, f"Approved By: {approved_by}", size=10)

    # Tutorial questions
    if data.get("tutorial_questions"):
        _heading2(doc, "Tutorial Questions with CO Mapping")
        _make_table(doc,
                    ["Q No", "Question", "CO"],
                    [[str(i+1), q.get("question_text",""), q.get("co_id","")]
                     for i, q in enumerate(data["tutorial_questions"])],
                    [1.0, 13.5, 2.0])

    # ── 7. Evaluation Plan & Marksheets ──────────────────────────────────────
    _section_title(doc, 7, "Evaluation plan with CO Mapping")
    eval_rows = data.get("eval_rows") or []
    if eval_rows:
        table_data = []
        for i, r in enumerate(eval_rows):
            sr   = _g(r, "ca_label", "label") or str(i + 1)
            comp = _g(r, "comp", "component", "name")
            units= _g(r, "unit_syllabus", "units", "syllabus", "unit")
            co   = _g(r, "co", "co_mapped")
            if isinstance(co, list):
                co = ", ".join(str(x) for x in co)
            marks= _g(r, "marks", "total_marks", "max_marks")
            wt   = _g(r, "weightage", "weight")
            date = _g(r, "date", "tentative_date")
            table_data.append([sr, comp, units, co, marks, wt, date])
        _make_table(doc,
                    ["Sr. No.", "Component", "Unit Syllabus", "CO", "Mark", "Weightage", "Tentative Date"],
                    table_data,
                    [1.0, 3.0, 4.5, 2.0, 1.2, 1.8, 3.0])
    else:
        _add_para(doc, "[Evaluation plan not yet generated. Use the Evaluation Plan page first.]",
                  color=(136, 136, 136))

    # ── 7b. Evaluation Components Details — Question Papers + Mark Sheets ─────
    _heading2(doc, "Evaluation Components Details")
    bloom_map = {1: "Remember", 2: "Understand", 3: "Apply",
                 4: "Analyse",  5: "Evaluate",   6: "Create"}

    ca_sheets_all = data.get("ca_sheets") or []
    student_map   = data.get("student_map") or {}
    students_list = data.get("students") or []

    for ca in ca_sheets_all:
        qp         = ca.get("qp") or []
        marks_data = ca.get("marks") or {}
        ca_label   = ca.get("ca_label", "")
        if not ca_label:
            continue

        # ── Question Paper ────────────────────────────────────────────────────
        if qp:
            _heading2(doc, f"{ca_label} — Question Paper")
            bl_map_for_ca = data.get("bloom_ai_map", {}).get(ca_label, {})
            qp_rows = []
            for q in qp:
                q_no   = q.get("q_no","")
                q_text = q.get("question_text","")
                marks  = q.get("marks","")
                co     = q.get("co_id","")
                bl_raw = q.get("bloom_level","")
                if bl_raw and isinstance(bl_raw, int):
                    bl_str = bloom_map.get(bl_raw, str(bl_raw))
                elif bl_raw:
                    bl_str = str(bl_raw)
                else:
                    bl_str = ""
                bl_str = bl_map_for_ca.get(str(q_no), bl_str) or bl_str
                qp_rows.append([q_no, q_text, marks, co, bl_str])
            _make_table(doc,
                        ["Q.No", "Question", "Marks", "CO", "Bloom's Level"],
                        qp_rows,
                        [1.0, 9.5, 1.2, 1.8, 2.0])

        # ── CA Announcement / Declaration ─────────────────────────────────────
        _heading2(doc, f"{ca_label} Result Declaration")

        if not qp:
            continue

        has_real_marks = any(
            any(float(v or 0) > 0 for v in mks.values())
            for mks in marks_data.values()
            if isinstance(mks, dict)
        )

        q_nos       = [q.get("q_no","") for q in qp]
        q_max_marks = [float(q.get("marks", 0) or 0) for q in qp]
        total_max   = sum(q_max_marks)
        num_students = len(students_list)

        if has_real_marks and num_students:
            q_avgs = []
            for q in qp:
                vals = [
                    float((marks_data.get(s["prn"]) or {}).get(str(q.get("q_no","")), 0) or 0)
                    for s in students_list
                ]
                q_avgs.append(f"{sum(vals)/len(vals):.1f}" if vals else "—")
            all_totals = []
            for s in students_list:
                mks = marks_data.get(s["prn"]) or {}
                all_totals.append(sum(float(mks.get(str(q.get("q_no","")),0) or 0) for q in qp))
            overall_avg = f"{sum(all_totals)/len(all_totals):.1f}" if all_totals else "—"
        else:
            q_avgs = ["0.0"] * len(qp)
            overall_avg = "0.0"

        # Mark sheet institutional header
        ms_tbl = doc.add_table(rows=1, cols=1)
        ms_tbl.style = "Table Grid"
        ms_cell = ms_tbl.rows[0].cells[0]
        _set_cell_bg(ms_cell, _NAVY)
        _set_cell_margins(ms_cell, top=80, bottom=80)
        p = ms_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "Department of AIML", bold=True, size=10, color=_WHITE)
        for ms_line in [
            f"Subject: {data.get('course_name','')} (Theory)",
            f"Batch: {batch}  Semester {data.get('semester','')}",
            f"{ca_label}: {int(total_max) if total_max else '?'} Marks",
        ]:
            p2 = ms_cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p2, ms_line, size=9, color=_WHITE)

        doc.add_paragraph()

        # Build CO sub-headers for each question
        q_co_headers = [q.get("co_id","") for q in qp]
        n = len(q_nos)
        col_w = [1.5, 4.5] + [round(8.0/max(n,1), 2)]*n + [1.5]
        num_cols = 2 + n + 1

        tbl2 = doc.add_table(rows=2 + (len(students_list) + 1), cols=num_cols)
        tbl2.style = "Table Grid"

        # Row 0: CO sub-header labels
        sr0 = tbl2.rows[0]
        for ci2 in range(num_cols):
            _set_cell_bg(sr0.cells[ci2], _NAVY)
            _set_cell_margins(sr0.cells[ci2])
            sr0.cells[ci2].width = Cm(col_w[ci2])
        for qi, co_id in enumerate(q_co_headers):
            p = sr0.cells[2+qi].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, co_id, bold=True, size=7, color=_WHITE)

        # Row 1: column labels
        sr1 = tbl2.rows[1]
        headers_main = (["Sr. No.", "Name of the Student"] +
                        [f"{q_nos[i]}\n(/{int(q_max_marks[i]) if q_max_marks[i] else '?'})"
                         for i in range(n)] +
                        ["Total"])
        for ci2, hdr_txt in enumerate(headers_main):
            cell = sr1.cells[ci2]
            cell.width = Cm(col_w[ci2])
            _set_cell_bg(cell, _NAVY)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, hdr_txt, bold=True, size=8, color=_WHITE)

        # Row 2: class average
        avg_tr = tbl2.rows[2]
        avg_vals = ["—", "Class Average"] + q_avgs + [overall_avg]
        for ci2, val in enumerate(avg_vals):
            cell = avg_tr.cells[ci2]
            cell.width = Cm(col_w[ci2])
            _set_cell_bg(cell, _GREEN)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci2 >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            _run(p, str(val or ""), bold=True, size=8, color=(0,100,0))

        # Data rows — students
        for ri, s in enumerate(students_list):
            prn  = s["prn"]
            name = s.get("name","")
            mks  = marks_data.get(prn) or {}
            row_vals = [str(ri+1), name]
            tot = 0.0
            for q in qp:
                v = float(mks.get(str(q.get("q_no","")), 0) or 0)
                row_vals.append(str(v) if has_real_marks and v else "")
                tot += v
            row_vals.append(f"{tot:.1f}" if has_real_marks and tot else "")
            bg = _WHITE if ri % 2 == 0 else _LIGHT
            tr = tbl2.rows[3 + ri]
            for ci2, val in enumerate(row_vals):
                cell = tr.cells[ci2]
                cell.width = Cm(col_w[ci2])
                _set_cell_bg(cell, bg)
                _set_cell_margins(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci2 >= 2 else WD_ALIGN_PARAGRAPH.LEFT
                _run(p, str(val or ""), size=8)

        doc.add_paragraph()

    # ── 7c. Final Marks Table ─────────────────────────────────────────────────
    if ca_sheets_all:
        _heading2(doc, "Final Marks Out of 30")
        final_header_tbl = doc.add_table(rows=1, cols=1)
        final_header_tbl.style = "Table Grid"
        fh_cell = final_header_tbl.rows[0].cells[0]
        _set_cell_bg(fh_cell, _NAVY)
        _set_cell_margins(fh_cell, top=80, bottom=80)
        p = fh_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, inst_name, bold=True, size=10, color=_WHITE)
        for fh_line in [
            f"Btech AIML / {batch} / {data.get('semester','')}",
            f"{data.get('course_name','')} / {data.get('course_code','')}",
        ]:
            pfh = fh_cell.add_paragraph()
            pfh.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(pfh, fh_line, size=9, color=_WHITE)
        doc.add_paragraph()

        ca_labels = [ca.get("ca_label","") for ca in ca_sheets_all if ca.get("ca_label")]
        ca_maxes  = [sum(float(q.get("marks",0) or 0) for q in (ca.get("qp") or []))
                     for ca in ca_sheets_all]
        grand_max = sum(ca_maxes)

        final_col_w = [1.0, 5.0, 2.5] + [round(4.0/max(len(ca_labels),1), 2)]*len(ca_labels) + [2.0, 2.0]
        final_headers = ["Sr. No.", "Name of the Student", "PRN"] + ca_labels + ["Total", "Final Marks"]

        final_rows = []
        for i, s in enumerate(students_list):
            prn  = s["prn"]
            name = s.get("name","")
            row  = [str(i+1), name, prn]
            grand = 0.0
            for ca in ca_sheets_all:
                qp2        = ca.get("qp") or []
                marks_data2 = ca.get("marks") or {}
                mks         = marks_data2.get(prn) or {}
                tot         = sum(float(mks.get(str(q.get("q_no","")),0) or 0) for q in qp2)
                row.append(f"{tot:.1f}" if tot else "")
                grand += tot
            row.append(f"{grand:.1f}" if grand else "")
            row.append(f"{grand:.2f}" if grand else "")
            final_rows.append(row)

        _make_table(doc, final_headers, final_rows, final_col_w)

    if not ca_sheets_all:
        _add_para(doc, "[No evaluation components found. Generate evaluation plan first.]",
                  color=(136, 136, 136))

    # ── 8. Slow & Advanced Learners ───────────────────────────────────────────
    _section_title(doc, 8, "List of Slow and Advanced learners and the action plans")

    _heading2(doc, "Advance Learners")
    advanced_list = data.get("advanced_learners_parsed") or []

    # FIX: parse plain-text advanced learners blob if computed list is empty
    if not advanced_list and data.get("advanced_learners"):
        raw_lines = [l.strip() for l in data["advanced_learners"].split("\n") if l.strip()]
        parsed = []
        for line in raw_lines:
            # Skip header/separator rows
            if line.startswith("-") or line.startswith("Student Name") or line.startswith("==="):
                continue
            parts = [p.strip() for p in line.split("|")]
            if len(parts) >= 2:
                name  = parts[0].strip()
                prn   = parts[1].strip() if len(parts) > 1 else ""
                extra = parts[2].strip() if len(parts) > 2 else ""
                if name:
                    parsed.append({"name": name, "prn": prn, "cgpa": extra})
            elif line.strip():
                parsed.append({"name": line.strip(), "prn": "", "cgpa": ""})
        advanced_list = parsed

    # Institutional header for advanced learners table
    adv_hdr_tbl = doc.add_table(rows=1, cols=1)
    adv_hdr_tbl.style = "Table Grid"
    adv_hdr_cell = adv_hdr_tbl.rows[0].cells[0]
    _set_cell_bg(adv_hdr_cell, _NAVY)
    _set_cell_margins(adv_hdr_cell, top=60, bottom=60)
    p = adv_hdr_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"AIML {batch} List of Students with CGPA more than 8.5 (Based on Result of previous Sem)",
         bold=True, size=9, color=_WHITE)
    doc.add_paragraph()

    if advanced_list:
        _make_table(doc,
                    ["PRN", "Name", "CGPA"],
                    [[s.get("prn",""), s.get("name",""), s.get("cgpa", s.get("marks",""))]
                     for s in advanced_list],
                    [3.0, 9.0, 4.5])
    else:
        _add_para(doc, "[Advanced learner list will appear once CA marks are entered.]",
                  color=(136, 136, 136))

    _heading2(doc, "Slow Learners")
    slow_list = data.get("slow_learners_parsed") or []

    # Institutional header for slow learners table
    sl_hdr_tbl = doc.add_table(rows=1, cols=1)
    sl_hdr_tbl.style = "Table Grid"
    sl_hdr_cell = sl_hdr_tbl.rows[0].cells[0]
    _set_cell_bg(sl_hdr_cell, _NAVY)
    _set_cell_margins(sl_hdr_cell, top=60, bottom=60)
    p = sl_hdr_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"AIML {batch} List of Students with CGPA less than 4.5 (Based on Result of previous Sem)",
         bold=True, size=9, color=_WHITE)
    doc.add_paragraph()

    if slow_list:
        _make_table(doc,
                    ["PRN", "Name", "CGPA", "SEM CGPA STATUS"],
                    [[s.get("prn",""), s.get("name",""), s.get("marks",""), "FAIL"]
                     for s in slow_list],
                    [3.0, 7.0, 2.0, 4.5])
    elif data.get("slow_learners"):
        lines = [l.strip() for l in data["slow_learners"].split("\n") if l.strip()]
        if lines:
            _make_table(doc,
                        ["PRN", "Name", "CGPA", "SEM CGPA STATUS"],
                        [["", l, "", ""] for l in lines],
                        [3.0, 7.0, 2.0, 4.5])
        else:
            _add_para(doc, "[Slow learner list will appear once CA marks are entered.]",
                      color=(136, 136, 136))
    else:
        _add_para(doc, "[Slow learner list will appear once CA marks are entered.]",
                  color=(136, 136, 136))

    # ── 9. CO Attainment (internal) ───────────────────────────────────────────
    _section_title(doc, 9, "CO Attainment of internal evaluation")

    co_attainment = data.get("co_attainment") or {}
    cos_list      = data.get("cos") or []

    # Institutional header
    att_hdr_tbl = doc.add_table(rows=1, cols=1)
    att_hdr_tbl.style = "Table Grid"
    att_hdr_cell = att_hdr_tbl.rows[0].cells[0]
    _set_cell_bg(att_hdr_cell, _NAVY)
    _set_cell_margins(att_hdr_cell, top=80, bottom=80)
    p = att_hdr_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Department of : Artificial Intelligence and Machine Learning",
         bold=True, size=10, color=_WHITE)
    for att_line in [
        "CO Attainment",
        f"Academic Year: {data.get('academic_year','')}    Batch: {batch}",
        f"Examination Season: {data.get('exam_season','APRIL 2025')}",
        f"Course Name: {data.get('course_name','')}    Course Code: {data.get('course_code','')}",
    ]:
        p2 = att_hdr_cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p2, att_line, size=9, color=_WHITE)
    doc.add_paragraph()

    if co_attainment or cos_list:
        ca_labels_att = [ca.get("ca_label","") for ca in (data.get("ca_sheets") or []) if ca.get("ca_label")]
        att_headers   = (["CO No / Weightage"] + ca_labels_att +
                         ["Internal\n100", "External\n0", "Final\n100", "Overall Att"])
        att_rows = []
        for co in (cos_list or [{"co_id": k} for k in co_attainment]):
            co_id = co.get("co_id","")
            pct   = co_attainment.get(co_id)
            level = 3 if pct and pct >= 70 else (2 if pct and pct >= 40 else 1)
            ca_vals = []
            for ca in (data.get("ca_sheets") or []):
                qp = ca.get("qp") or []
                co_qs = [q for q in qp if q.get("co_id") == co_id]
                ca_vals.append(str(level) if co_qs else "")
            internal_val = f"{pct:.2f}" if pct else ""
            att_rows.append([co_id] + ca_vals + [internal_val, "", internal_val,
                                                  f"{pct:.2f}" if pct else ""])
        col_w_att = ([2.0] + [round(9.0/max(len(ca_labels_att),1), 2)]*len(ca_labels_att) +
                     [2.0, 2.0, 2.0, 2.0])
        _make_table(doc, att_headers, att_rows, col_w_att)

        # Footer note
        doc.add_paragraph()
        note_tbl = doc.add_table(rows=1, cols=3)
        note_tbl.style = "Table Grid"
        for ci2, (label, val) in enumerate([
            ("External", "60%  OR  100%  OR  Nil"),
            ("Internal", "40%  |   NIL   |  100%"),
            ("",          "Both | Only ESE | Only CIE"),
        ]):
            cell = note_tbl.rows[0].cells[ci2]
            _set_cell_bg(cell, _LGRAY)
            _set_cell_margins(cell, top=60, bottom=60)
            p = cell.paragraphs[0]
            if label:
                r1 = p.add_run(f"{label}: ")
                r1.bold = True
                r1.font.size = Pt(8)
            p.add_run(val).font.size = Pt(8)
    else:
        _add_para(doc, "[CO attainment will appear here once marks are entered.]",
                  color=(136, 136, 136))

    doc.add_paragraph()

    # PO-level attainment matrix
    if co_attainment and cos_list:
        _heading2(doc, "PO / PSO Attainment")
        po_ids_att = [f"PO{i}" for i in range(1, 13)] + ["PSO1", "PSO2"]

        # Use standard mapping if DB values are all-1s
        def _effective_mapping(co_id):
            if use_standard_mapping:
                return _STANDARD_CO_PO.get(co_id, {})
            return (co_po_matrix or {}).get(co_id) or {}

        # Build CO attainment levels (convert % to 1/2/3 scale)
        # Formula from reference doc: ≥70% → 3, ≥40% → 2, else → 1
        def _co_level(co_id):
            pct = co_attainment.get(co_id)
            if not pct:
                return None
            return 3 if pct >= 70 else (2 if pct >= 40 else 1)

        po_att_rows = []
        for co in cos_list:
            co_id   = co.get("co_id","")
            mapping = _effective_mapping(co_id)
            def _gm(pid, _m=mapping):
                v = _m.get(pid) or _m.get(pid.replace("PO","PO ")) or _m.get(pid.replace("PO ","PO"))
                return str(v) if v and str(v) not in ("0","") else ""
            po_att_rows.append([_gm(pid) for pid in po_ids_att])

        # PO attainment = weighted avg of (CO_level × CO_PO_strength) for mapped COs
        # weighted by CO_PO_strength, per NBA/NAAC standard formula
        po_attainment_vals = []
        for idx, pid in enumerate(po_ids_att):
            numerator, denominator = 0.0, 0.0
            for co, row_vals in zip(cos_list, po_att_rows):
                co_id  = co.get("co_id","")
                level  = _co_level(co_id)
                m_str  = row_vals[idx]
                if m_str and level is not None:
                    try:
                        strength = float(m_str)
                        if strength > 0:
                            numerator   += level * strength
                            denominator += strength
                    except (ValueError, TypeError):
                        pass
            po_attainment_vals.append(f"{numerator/denominator:.2f}" if denominator > 0 else "-")

        # Row 1: CO-PO mapping strengths (from database / AI / standard)
        att2_headers = ["CO"] + po_ids_att
        att2_rows    = [[co.get("co_id","")] + row_vals for co, row_vals in zip(cos_list, po_att_rows)]

        # Row 2: CO attainment levels
        level_row = ["CO Att. Level"]
        for co in cos_list:
            lv = _co_level(co.get("co_id",""))
            level_row.append(str(lv) if lv else "-")
        # pad level row to match columns (one value per CO, not per PO — different shape)
        # Instead: show CO attainment % in a separate summary row
        co_pct_row = ["CO Att %"]
        for co in cos_list:
            pct = co_attainment.get(co.get("co_id",""))
            co_pct_row.append(f"{pct:.1f}%" if pct else "-")

        att2_rows.append(["PO Attainment"] + po_attainment_vals)

        n_po     = len(po_ids_att)
        col_w_po = [1.5] + [round(14.5/max(n_po,1), 2)]*n_po
        _make_table(doc, att2_headers, att2_rows, col_w_po)

        # CO attainment summary table
        doc.add_paragraph()
        _heading2(doc, "CO Attainment Summary")
        co_summary_headers = ["CO"] + [co.get("co_id","") for co in cos_list]
        pct_vals = [f"{co_attainment.get(co.get('co_id',''), 0):.1f}%" if co_attainment.get(co.get('co_id','')) else "-"
                    for co in cos_list]
        level_vals = [str(_co_level(co.get("co_id",""))) if _co_level(co.get("co_id","")) else "-"
                      for co in cos_list]
        co_sum_rows = [["Attainment %"] + pct_vals, ["Level (1/2/3)"] + level_vals]
        col_w_sum = [2.5] + [round(12.5/max(len(cos_list),1), 2)]*len(cos_list)
        _make_table(doc, co_summary_headers, co_sum_rows, col_w_sum)

    # ── 10. Activity Reports ──────────────────────────────────────────────────
    _section_title(doc, 10, "The reports of the activities planned and conducted")
    activity_reports = data.get("activity_reports") or ""
    if activity_reports.strip():
        try:
            reports = json.loads(activity_reports)
            if isinstance(reports, list):
                _heading2(doc, "Best Practice and Innovative Activities-")
                for i, rpt in enumerate(reports):
                    if isinstance(rpt, dict):
                        _add_para(doc, f"{i+1}.\t{rpt.get('title','')}", bold=True, size=11)
                        details_map = {
                            "conduction_date":   "Conduction Date",
                            "time_duration":     "Time (Duration)",
                            "total_hours":       "Total No. of Hours",
                            "venue":             "Venue",
                            "attended_by":       "Attended by (Batch with Branch)",
                            "students_attended": "No. Of Student attended the session",
                            "staff_attended":    "No. Of Staff attended the session",
                            "arranged_by":       "Arranged by",
                        }
                        for field, label in details_map.items():
                            val = rpt.get(field, "")
                            if val:
                                p = doc.add_paragraph()
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.space_after  = Pt(2)
                                r1 = p.add_run(f"{label}   - ")
                                r1.bold = True
                                r1.font.size = Pt(10)
                                p.add_run(str(val)).font.size = Pt(10)
                        if rpt.get("speaker_name"):
                            _heading2(doc, "About Speaker")
                            for sk, sv in [("speaker_name","Name"), ("company","Company Name"),
                                           ("designation","Designation"), ("contact","Contact Details")]:
                                if rpt.get(sk):
                                    p = doc.add_paragraph()
                                    r1 = p.add_run(f"{sv} \u2013 ")
                                    r1.bold = True
                                    r1.font.size = Pt(10)
                                    p.add_run(str(rpt[sk])).font.size = Pt(10)
                        if rpt.get("report"):
                            _heading2(doc, "Event Report in brief:")
                            _add_para(doc, rpt["report"])
                        if rpt.get("topics"):
                            _heading2(doc, "Topics Covered")
                            for t in (rpt["topics"] if isinstance(rpt["topics"], list) else [rpt["topics"]]):
                                p = doc.add_paragraph(style="List Bullet")
                                _run(p, str(t), size=10)
                        if rpt.get("outcomes"):
                            _heading2(doc, "Outcomes")
                            outcomes = rpt["outcomes"] if isinstance(rpt["outcomes"], list) else [rpt["outcomes"]]
                            out_tbl = doc.add_table(rows=len(outcomes), cols=1)
                            out_tbl.style = "Table Grid"
                            for oi, outcome in enumerate(outcomes):
                                cell = out_tbl.rows[oi].cells[0]
                                _set_cell_bg(cell, _WHITE if oi % 2 == 0 else _LIGHT)
                                _set_cell_margins(cell)
                                _run(cell.paragraphs[0], str(outcome), size=9)
                        if rpt.get("feedback_link"):
                            _heading2(doc, "Feedback")
                            p = doc.add_paragraph()
                            r = p.add_run(rpt["feedback_link"])
                            r.font.color.rgb = _rgb((5, 99, 193))
                            r.underline = True
                            r.font.size = Pt(10)
                    else:
                        _add_para(doc, f"{i+1}. {rpt}")
            else:
                raise ValueError("not a list")
        except (json.JSONDecodeError, ValueError):
            lines = [l.strip() for l in activity_reports.split("\n") if l.strip()]
            _heading2(doc, "Best Practice and Innovative Activities-")
            for i, line in enumerate(lines):
                _add_para(doc, f"{i+1}.\t{line}", space_before=2, space_after=2)
    else:
        _add_para(doc, "[Activity reports not yet entered. Add them in the Course File section.]",
                  color=(136, 136, 136))

    # ── 11. Learning Material ─────────────────────────────────────────────────
    _section_title(doc, 11, "Learning Material.")

    # Parse the raw learning_material_links text into structured categories
    def _parse_learning_materials(raw: str):
        """
        Parse newline-separated learning material entries into typed buckets.
        Detects: Textbook, Journal, MOOC/Coursera/NPTEL platform, research articles, web links.
        """
        import re as _re
        tb, web, jour, mooc, art = [], [], [], [], []
        for line in raw.split("\n"):
            line = line.strip()
            if not line:
                continue
            lo = line.lower()
            # Extract URL if present inline
            url_match = _re.search(r"https?://\S+", line)
            url  = url_match.group(0) if url_match else ""
            text = line.replace(url, "").strip(" —:-") if url else line

            if lo.startswith("textbook"):
                # "Textbook: Title — Author, Publisher"  or  "Textbook: Title"
                body = _re.sub(r"^textbook\s*[:\-]?\s*", "", line, flags=_re.IGNORECASE).strip()
                parts = _re.split(r"\s*[—–-]{1,2}\s*", body, maxsplit=1)
                title = parts[0].strip()
                rest  = parts[1].strip() if len(parts) > 1 else ""
                # Try to split author, publisher by last comma
                if "," in rest:
                    idx = rest.rfind(",")
                    author    = rest[:idx].strip()
                    publisher = rest[idx+1:].strip()
                else:
                    author, publisher = rest, ""
                tb.append({"title": title, "author": author, "publisher": publisher})

            elif any(k in lo for k in ["journal:", "journal of ", "ieee ", "transactions on", "pattern recognition"]):
                title = _re.sub(r"^journal\s*[:\-]?\s*", "", line, flags=_re.IGNORECASE).strip()
                jour.append({"title": title})

            elif any(k in lo for k in ["coursera", "nptel", "edx", "udemy", "swayam", "mit-ocw", "mooc"]):
                platform = ("Coursera" if "coursera" in lo else
                            "NPTEL"    if "nptel"    in lo else
                            "edX"      if "edx"      in lo else
                            "Udemy"    if "udemy"    in lo else
                            "SWAYAM"   if "swayam"   in lo else "Online")
                mooc.append({"title": text or line, "platform": platform, "duration": "", "certificate": "Y", "url": url})

            elif any(k in lo for k in ["arxiv", "et al", "generative", "variational", "tutorial on", "review"]):
                art.append({"title": text or line, "url": url})

            else:
                # Generic web link or reference
                web.append({"title": text or line, "url": url, "unit": ""})

        return tb, web, jour, mooc, art

    # Prefer structured study_materials; fall back to parsing raw text
    mat11  = data.get("study_materials") or {}
    tb11   = mat11.get("textbooks") or []
    ref11  = mat11.get("reference_books") or mat11.get("references") or []
    web11  = mat11.get("web_links") or mat11.get("web") or []
    jour11 = mat11.get("journals") or []
    mooc11 = mat11.get("moocs") or []
    art11  = mat11.get("research_articles") or []

    raw_links = (data.get("learning_material_links") or "").strip()
    if raw_links and not any([tb11, ref11, web11, jour11, mooc11, art11]):
        _ptb, _pweb, _pjour, _pmooc, _part = _parse_learning_materials(raw_links)
        tb11   = _ptb
        web11  = _pweb
        jour11 = _pjour
        mooc11 = _pmooc
        art11  = _part

    if tb11 or ref11:
        _heading2(doc, "Textbooks & Reference books/ Beyond Gaps")
        _make_table(doc, ["Book", "Author", "Publisher"],
                    [[b.get("title", str(b) if not isinstance(b,dict) else ""),
                      b.get("author","") if isinstance(b,dict) else "",
                      b.get("publisher","") if isinstance(b,dict) else ""]
                     for b in (tb11 + ref11)],
                    [7.0, 4.0, 5.0])

    if web11:
        _heading2(doc, "Web-Links for Online Notes/ YouTube/NPTEL Videos/Blogs etc")
        _make_table(doc, ["Sr. No.", "Web Link", "Module"],
                    [[str(i+1),
                      w.get("title", w.get("url","") if isinstance(w,dict) else str(w)),
                      w.get("unit", w.get("module","")) if isinstance(w,dict) else ""]
                     for i, w in enumerate(web11)],
                    [1.0, 10.0, 5.0])

    if jour11:
        _heading2(doc, "Names of Magazines, Journals, E-journals")
        _make_table(doc, ["Sr.No.", "Journal"],
                    [[str(i+1),
                      j.get("title","") if isinstance(j,dict) else str(j)]
                     for i, j in enumerate(jour11)],
                    [1.0, 15.0])

    if mooc11:
        _heading2(doc, "Recommended MOOC Courses like Coursera / NPTEL / MIT-OCW / edX etc")
        _make_table(doc,
                    ["S.No.", "MOOC Course Link", "Course conducted by", "Course Duration", "Certificate (Y / N)"],
                    [[str(i+1),
                      m.get("title", m.get("url","") if isinstance(m,dict) else str(m)),
                      m.get("platform", m.get("conducted_by","")) if isinstance(m,dict) else "",
                      m.get("duration","") if isinstance(m,dict) else "",
                      m.get("certificate","Y") if isinstance(m,dict) else "Y"]
                     for i, m in enumerate(mooc11)],
                    [1.0, 6.0, 3.0, 2.5, 2.0])

    if art11:
        _heading2(doc, "List of Research Articles")
        _make_table(doc, ["S.No.", "Research Article Title", "Web Link"],
                    [[str(i+1),
                      a.get("title","") if isinstance(a,dict) else str(a),
                      a.get("url","") if isinstance(a,dict) else ""]
                     for i, a in enumerate(art11)],
                    [1.0, 9.0, 6.0])

    if not any([tb11, ref11, web11, jour11, mooc11, art11]):
        _add_para(doc, "[Learning material not yet entered.]", color=(136, 136, 136))

    # ── 12. Question Bank ─────────────────────────────────────────────────────
    _section_title(doc, 12, "Question Bank")

    # Merge DB questions + all CA/ESE sheet questions so bank is never empty
    # Dedup key = (question_text, co_id, source) so same text under different CO is kept
    questions = list(data.get("questions") or [])
    seen = {(q.get("question_text","").strip().lower(), q.get("co_id",""), "") for q in questions}
    for ca in (data.get("ca_sheets") or []):
        ca_label = ca.get("ca_label", "")
        for q in (ca.get("qp") or []):
            qt  = (q.get("question_text") or "").strip()
            co  = q.get("co_id", "")
            key = (qt.lower(), co, ca_label)
            if qt and key not in seen:
                questions.append({
                    "question_text": qt,
                    "co_id":         co,
                    "bloom_level":   q.get("bloom_level", ""),
                    "marks":         q.get("marks", ""),
                    "unit_no":       q.get("unit_no", ""),
                    "source":        ca_label,
                })
                seen.add(key)

    if questions:
        from collections import defaultdict as _dd2
        # Group by CO
        by_co = _dd2(list)
        for q in questions:
            co_key = q.get("co_id") or "General"
            by_co[co_key].append(q)

        for co_label in sorted(by_co.keys()):
            co_qs = by_co[co_label]
            # CO sub-heading — teal banner
            doc.add_paragraph()
            tbl_co = doc.add_table(rows=1, cols=1)
            tbl_co.allow_autofit = False
            hcell = tbl_co.rows[0].cells[0]
            hcell.width = Cm(16)
            _set_cell_bg(hcell, _TEAL)
            hp = hcell.paragraphs[0]
            hp.paragraph_format.space_before = Pt(3)
            hp.paragraph_format.space_after  = Pt(3)
            _run(hp, f"  {co_label}", bold=True, size=11, color=_WHITE)

            # Question table: Q.No | Question | Bloom's | Marks | Source
            q_rows = []
            for i, q in enumerate(co_qs):
                bloom = q.get("bloom_level") or ""
                marks = str(q.get("marks") or "")
                src   = q.get("source") or ""
                q_rows.append([
                    str(i + 1),
                    q.get("question_text", ""),
                    bloom,
                    marks,
                    src,
                ])

            tbl = doc.add_table(rows=1 + len(q_rows), cols=5)
            tbl.allow_autofit = False
            tbl.style = "Table Grid"
            col_widths = [Cm(0.8), Cm(10.0), Cm(2.2), Cm(1.4), Cm(1.6)]
            headers = ["#", "Question", "Bloom's Level", "Marks", "Source"]

            # Header row
            for ci, (hdr, cw) in enumerate(zip(headers, col_widths)):
                cell = tbl.rows[0].cells[ci]
                cell.width = cw
                _set_cell_bg(cell, _NAVY)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _run(p, hdr, bold=True, size=9, color=_WHITE)

            # Data rows — alternating mint / white, generous padding
            for ri, row_vals in enumerate(q_rows):
                bg = _MINT if ri % 2 == 0 else _WHITE
                tr = tbl.rows[ri + 1]
                for ci, (val, cw) in enumerate(zip(row_vals, col_widths)):
                    cell = tr.cells[ci]
                    cell.width = cw
                    _set_cell_bg(cell, bg)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after  = Pt(4)
                    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                   if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    _run(p, str(val), bold=(ci == 0), size=9)

            doc.add_paragraph()
    else:
        _add_para(doc, "[Question bank is empty — add questions via the Question Bank page or upload CA sheets.]",
                  color=(136, 136, 136))

    # ── 13. Attendance ────────────────────────────────────────────────────────
    _section_title(doc, 13, "Compiled Attendance")
    if data.get("attendance_links"):
        links = [l.strip() for l in data["attendance_links"].split("\n") if l.strip()]
        div_labels = (["Division A", "Division B", "Division C"] +
                      [f"Division {chr(68+i)}" for i in range(10)])
        for i, link in enumerate(links):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            label = div_labels[i] if i < len(div_labels) else f"Division {i+1}"
            r1 = p.add_run(f"{label}  ")
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(link)
            r2.font.color.rgb = _rgb((5, 99, 193))
            r2.underline = True
            r2.font.size = Pt(10)
    else:
        _add_para(doc, "[Attendance links not yet entered. Add them in the Course File section.]",
                  color=(136, 136, 136))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Service class
# ─────────────────────────────────────────────────────────────────────────────

class CourseFileService:
    def __init__(self, db: AsyncSession, current_user=None):
        self.db = db
        self.current_user = current_user

    @staticmethod
    def get_filepath(course_id: int) -> str:
        storage = get_storage()
        p = storage.get_path(_CATEGORY, f"course_file_{course_id}.docx")
        return str(p) if p else str(
            get_storage()._dir(_CATEGORY) / f"course_file_{course_id}.docx"
        )

    async def _get_students(self, course_id: int):
        try:
            result = await self.db.execute(
                text("SELECT prn, name, section FROM students WHERE course_id=:cid ORDER BY section, name"),
                {"cid": course_id}
            )
            return [{"prn": r[0], "name": r[1], "section": r[2]} for r in result.fetchall()]
        except Exception:
            return []

    async def _get_session_rows(self, course_id: int):
        from backend.database.models import SessionPlanRow
        result = await self.db.execute(
            select(SessionPlanRow).where(SessionPlanRow.course_id == course_id)
        )
        row = result.scalar_one_or_none()
        return row.rows if row else []

    async def _get_eval_rows(self, course_id: int):
        from backend.database.models import EvalPlanRow
        result = await self.db.execute(
            select(EvalPlanRow).where(EvalPlanRow.course_id == course_id)
        )
        row = result.scalar_one_or_none()
        return row.rows if row else []

    async def _get_ca_sheets(self, course_id: int):
        from backend.database.models import CASheet
        result = await self.db.execute(
            select(CASheet).where(CASheet.course_id == course_id)
        )
        sheets = result.scalars().all()
        return [{"ca_label": s.ca_label, "qp": s.qp, "marks": s.marks} for s in sheets]

    async def _get_questions(self, course_id: int):
        from backend.database.models import Question
        result = await self.db.execute(
            select(Question).where(Question.course_id == course_id)
        )
        qs = result.scalars().all()
        return [{"question_text": q.question_text, "co_id": q.co_id,
                 "bloom_level": q.bloom_level, "marks": q.marks,
                 "unit_no": getattr(q, "unit_no", None)} for q in qs]

    async def _get_timetable(self) -> dict:
        """Pull the uploaded timetable from the current user DB record (dashboard upload)."""
        try:
            if self.current_user and getattr(self.current_user, "timetable_json", None):
                return json.loads(self.current_user.timetable_json)
        except Exception as e:
            logger.warning(f"Could not read timetable from user record: {e}")
        # Fallback: legacy file-based storage
        try:
            storage = get_storage()
            path = storage.get_path("timetables", "current_timetable.json")
            if path and Path(path).exists():
                return json.loads(Path(path).read_text(encoding="utf-8"))
        except Exception as e:
            logger.warning(f"Could not read timetable from file storage: {e}")
        return {}

    async def _get_study_materials(self, course_id: int) -> dict:
        """Try to pull study materials from the session plan materials endpoint data."""
        try:
            from backend.database.models import SessionPlanRow
            result = await self.db.execute(
                select(SessionPlanRow).where(SessionPlanRow.course_id == course_id)
            )
            sp_row = result.scalar_one_or_none()
            if not sp_row or not sp_row.rows:
                return {}

            rows = sp_row.rows
            cols = sp_row.cols or []

            textbooks, ref_books, web_links, journals, moocs, research_articles = [], [], [], [], [], []

            for col in cols:
                label = col.get("label", "").lower()
                key   = col.get("key", "")
                if not key:
                    continue
                if any(k in label for k in ["textbook", "text book"]):
                    for r in rows:
                        v = r.get(key, "")
                        if v and v not in [t.get("title","") for t in textbooks]:
                            textbooks.append({"title": v, "author": "", "publisher": ""})
                elif any(k in label for k in ["reference", "ref book"]):
                    for r in rows:
                        v = r.get(key, "")
                        if v and v not in [t.get("title","") for t in ref_books]:
                            ref_books.append({"title": v, "author": "", "publisher": ""})
                elif any(k in label for k in ["web", "link", "nptel", "url", "online", "swayam"]):
                    for r in rows:
                        v = r.get(key, "")
                        if v and v not in [w.get("title","") for w in web_links]:
                            web_links.append({"title": v, "unit": r.get("unit", ""), "url": ""})
                elif any(k in label for k in ["research", "article", "paper", "classic"]):
                    for r in rows:
                        v = r.get(key, "")
                        if v and v not in [a.get("title","") for a in research_articles]:
                            research_articles.append({"title": v, "url": ""})
                elif any(k in label for k in ["journal", "magazine"]):
                    for r in rows:
                        v = r.get(key, "")
                        if v and v not in [j.get("title","") for j in journals]:
                            journals.append({"title": v, "url": ""})
                elif any(k in label for k in ["mooc", "course", "coursera", "swayam", "edx"]):
                    for r in rows:
                        v = r.get(key, "")
                        if v and v not in [m.get("title","") for m in moocs]:
                            moocs.append({"title": v, "platform": "", "duration": "", "certificate": "Y"})

            return {
                "textbooks": textbooks,
                "reference_books": ref_books,
                "web_links": web_links,
                "journals": journals,
                "moocs": moocs,
                "research_articles": research_articles,
            }
        except Exception as e:
            logger.warning(f"Could not read study materials: {e}")
            return {}

    async def _get_co_attainment(self, course_id, students, ca_sheets, cos):
        attainment = {}
        for co in cos:
            cid = co["co_id"]
            total_pct, count = 0, 0
            for sheet in ca_sheets:
                qp = [q for q in (sheet.get("qp") or []) if q.get("co_id") == cid]
                max_marks = sum(float(q.get("marks", 0)) for q in qp)
                if not max_marks or not students:
                    continue
                marks_data = sheet.get("marks") or {}
                has_real = any(
                    any(float(v or 0) > 0 for v in (mks or {}).values())
                    for mks in marks_data.values() if isinstance(mks, dict)
                )
                if not has_real:
                    continue
                passed = sum(
                    1 for s in students
                    if sum(float((marks_data.get(s["prn"]) or {}).get(str(q.get("q_no","")), 0))
                           for q in qp) / max_marks * 100 >= 60
                )
                total_pct += (passed / len(students)) * 100
                count += 1
            attainment[cid] = round(total_pct / count, 1) if count else None
        return {k: v for k, v in attainment.items() if v is not None}

    async def _get_slow_advanced(self, course_id, students, ca_sheets, cos):
        if not students or not ca_sheets:
            return [], []
        totals = {s["prn"]: 0.0 for s in students}
        maxes  = {s["prn"]: 0.0 for s in students}
        has_any_marks = False
        for sheet in ca_sheets:
            qp = sheet.get("qp") or []
            marks_data = sheet.get("marks") or {}
            total_marks = sum(float(q.get("marks", 0)) for q in qp)
            if not total_marks:
                continue
            for s in students:
                obtained = sum(float((marks_data.get(s["prn"]) or {}).get(str(q.get("q_no","")), 0))
                               for q in qp)
                if obtained > 0:
                    has_any_marks = True
                totals[s["prn"]] += obtained
                maxes[s["prn"]]  += total_marks

        if not has_any_marks:
            return [], []

        scored = []
        for s in students:
            mx  = maxes.get(s["prn"], 0)
            tot = totals.get(s["prn"], 0)
            pct = (tot / mx * 100) if mx else 0
            scored.append({"prn": s["prn"], "name": s["name"],
                           "marks": f"{tot:.1f}/{mx:.0f}", "pct": pct})
        scored.sort(key=lambda x: x["pct"])
        return [s for s in scored if s["pct"] < 40], [s for s in scored if s["pct"] >= 75]

    async def _get_extra(self, course_id: int):
        from backend.database.models import CourseFileExtra
        result = await self.db.execute(
            select(CourseFileExtra).where(CourseFileExtra.course_id == course_id)
        )
        extra = result.scalar_one_or_none()
        return extra.to_dict() if extra else {}

    async def _get_attachments(self, course_id: int):
        from backend.database.models import CourseFileAttachment
        result = await self.db.execute(
            select(CourseFileAttachment)
            .where(CourseFileAttachment.course_id == course_id)
            .order_by(CourseFileAttachment.section_no, CourseFileAttachment.uploaded_at)
        )
        return [a.to_dict() for a in result.scalars().all()]

    def _extract_syllabus_from_session(self, session_rows):
        units = {}
        for row in session_rows:
            unit_no = _g(row, "unit", "unit_no", "unitNo", "unit_number")
            topic   = _g(row, "topic", "points_to_cover", "pointsToCover", "content")
            if not unit_no:
                continue
            key = str(unit_no)
            if key not in units:
                units[key] = {"unit_number": unit_no,
                              "unit_title": row.get("unit_title", f"Unit {unit_no}"),
                              "topics": []}
            if topic and topic not in units[key]["topics"]:
                units[key]["topics"].append(topic)
        return list(units.values())

    def _extract_tutorial_questions(self, questions, max_per_co=5):
        if not questions:
            return []
        by_co = {}
        for q in questions:
            co = q.get("co_id") or "General"
            by_co.setdefault(co, []).append(q)
        result = []
        for co, qs in by_co.items():
            result.extend(qs[:max_per_co])
        return result

    async def _ai_bloom_map(self, ca_sheets: list) -> dict:
        """
        For each CA sheet, AI-classify any questions missing a bloom level.
        Returns {ca_label: {q_no: bloom_level_str}}.
        Falls back gracefully if LLM unavailable.
        """
        import json as _json
        from backend.core.llm import get_llm_response

        BLOOM_NAMES = ["Remember", "Understand", "Apply", "Analyse", "Evaluate", "Create"]
        result = {}

        for ca in (ca_sheets or []):
            qp = ca.get("qp") or []
            ca_label = ca.get("ca_label", "")
            if not qp:
                continue

            needs_ai = [
                q for q in qp
                if not q.get("bloom_level") or q.get("bloom_level") == 0
            ]
            if not needs_ai:
                continue

            lines = []
            for q in needs_ai:
                q_no = str(q.get("q_no", "?"))
                text = str(q.get("question_text", ""))[:120]
                lines.append("  Q" + q_no + ". " + text)
            q_list = "\n".join(lines)

            prompt = (
                "You are an expert in Bloom's Taxonomy for engineering education. "
                "Classify each question below into exactly ONE Bloom level from: "
                "Remember, Understand, Apply, Analyse, Evaluate, Create. "
                "Return a JSON object mapping Q-number (e.g. \"Q1\") to Bloom level string. "
                "Respond ONLY with the JSON, no markdown fences, no explanation.\n\n"
                "Questions:\n" + q_list
            )

            try:
                raw = await get_llm_response(prompt)
                raw = raw.strip()
                if raw.startswith("```"):
                    raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
                mapping = _json.loads(raw)
                clean = {}
                for k, v in mapping.items():
                    v_str = str(v).strip()
                    matched = next(
                        (b for b in BLOOM_NAMES if b.lower() in v_str.lower()), None
                    )
                    if matched:
                        key = str(k).replace("Q", "").strip()
                        clean[key] = matched
                result[ca_label] = clean
            except Exception as e:
                logger.warning("AI bloom mapping failed for %s: %s", ca_label, e)

        return result

    async def generate(self, course_id: int) -> dict:
        course_svc = CourseService(self.db)
        course     = await course_svc.get_course(course_id)
        students   = await self._get_students(course_id)
        session_rows = await self._get_session_rows(course_id)
        eval_rows    = await self._get_eval_rows(course_id)
        ca_sheets    = await self._get_ca_sheets(course_id)
        questions    = await self._get_questions(course_id)
        extra        = await self._get_extra(course_id)
        timetable    = await self._get_timetable()
        study_mats   = await self._get_study_materials(course_id)
        cos          = course.cos

        co_attainment              = await self._get_co_attainment(course_id, students, ca_sheets, cos)
        slow_list, advanced_list   = await self._get_slow_advanced(course_id, students, ca_sheets, cos)
        syllabus_units             = self._extract_syllabus_from_session(session_rows)
        tutorial_qs                = self._extract_tutorial_questions(questions)

        data = {
            "institution_name":    extra.get("institution_name", "Symbiosis Institute of Technology"),
            "institution_address": extra.get("institution_address", "SIU Pune 412115, Maharashtra, India"),
            "course_name":    course.course_name,
            "course_code":    course.course_code,
            "department":     course.department,
            "faculty_name":   course.faculty_name,
            "hod_name":       extra.get("hod_name", ""),
            "semester":       course.semester,
            "academic_year":  course.academic_year,
            "exam_season":    extra.get("exam_season", "APRIL 2025"),
            "credits":        course.credits,
            "batch":          extra.get("batch", ""),
            "cos":            cos,
            "pos":            course.pos,
            "co_po_matrix":   course.co_po_matrix,
            "co_po_justification": extra.get("co_po_justification", ""),
            "vision_text":    extra.get("vision_text", ""),
            "mission_text":   extra.get("mission_text", ""),
            "syllabus_units": syllabus_units,
            "prev_co_attainment": extra.get("prev_co_attainment", ""),
            "action_plan":    extra.get("action_plan", ""),
            "session_rows":   session_rows,
            "tutorial_questions": tutorial_qs,
            "eval_rows":      eval_rows,
            "ca_sheets":      ca_sheets,
            "student_map":    {s["prn"]: s["name"] for s in students},
            "slow_learners":  extra.get("slow_learners", ""),
            "advanced_learners": extra.get("advanced_learners", ""),
            "slow_learners_parsed":     slow_list,
            "advanced_learners_parsed": advanced_list,
            "students":       students,
            "timetable":      timetable,
            "study_materials": study_mats,
            "attachments":    await self._get_attachments(course_id),
            "co_attainment":  co_attainment,
            "activity_reports": extra.get("activity_reports", ""),
            "learning_material_links": extra.get("learning_material_links", ""),
            "questions":      questions,
            "attendance_links": extra.get("attendance_links", ""),
        }

        data["bloom_ai_map"] = await self._ai_bloom_map(ca_sheets)

        docx_bytes = _build_docx(data)

        _storage  = get_storage()
        _filename = f"course_file_{course_id}.docx"
        _storage.save(_CATEGORY, _filename, docx_bytes)

        filepath = str(_storage.get_path(_CATEGORY, _filename))
        logger.info(f"Course file saved -> {filepath}")

        return {
            "course_id":   course_id,
            "course_name": course.course_name,
            "filename":    _filename,
            "download_url": f"/course-file/download/{course_id}",
            "sections_with_data": {
                "session_plan":     len(session_rows) > 0,
                "evaluation_plan":  len(eval_rows) > 0,
                "ca_marks":         len(ca_sheets) > 0,
                "question_bank":    len(questions) > 0,
                "vision_mission":   bool(extra.get("vision_text")),
                "timetable":        bool(timetable),
                "study_materials":  bool(any(study_mats.values())),
            },
        }

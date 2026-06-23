# backend/services/attainment_service.py
"""
AttainmentService
-----------------
1. save_marks()   — bulk-upsert student marks for a course
2. calculate()    — compute CO attainment % and PO attainment, return dict
3. generate_report() — write a formatted Word document (.docx)

Attainment formula (standard OBE practice):
  CO attainment % = (students scoring >= threshold%) / total_students * 100
  Default threshold = 60% of max marks for that CO's components.

Output folder: generated_docs/attainment_reports/
"""

import json
import os
from pathlib import Path
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete

from backend.core.config import GEMINI_API_KEY
from backend.core.exceptions import OBEException
from backend.core.logger import get_logger
from backend.database.models import Course, COAttainment
from backend.services.course_service import CourseService

logger = get_logger(__name__)

from backend.core.storage import get_storage
_CATEGORY = "attainment_reports"

ATTAINMENT_THRESHOLD_PCT = 60   # students must score >= 60% of max CO marks


class AttainmentService:

    def __init__(self, db: AsyncSession):
        self.db = db

    # ------------------------------------------------------------------
    # Static helper
    # ------------------------------------------------------------------

    @staticmethod
    def get_filepath(course_id: int) -> str:
        storage = get_storage()
        p = storage.get_path(_CATEGORY, f"attainment_report_{course_id}.docx")
        return str(p) if p else str(get_storage()._dir(_CATEGORY) / f"attainment_report_{course_id}.docx")

    # ------------------------------------------------------------------
    # 1. Save / replace student marks
    # ------------------------------------------------------------------

    async def save_marks(self, course_id: int, students: list) -> dict:
        """
        Accepts a list of student mark dicts, replaces all existing records
        for this course, and commits.

        students = [
          {
            "student_id": "USN001",
            "student_name": "Alice",
            "marks": {
              "CO1": {"Quiz": 8, "Unit Test": 18},
              "CO2": {"Quiz": 7, "Unit Test": 15}
            }
          }, ...
        ]
        """
        # Verify course exists
        course_svc = CourseService(self.db)
        await course_svc.get_course(course_id)   # raises 404 if missing

        # Delete all existing records for this course
        await self.db.execute(
            delete(COAttainment).where(COAttainment.course_id == course_id)
        )
        await self.db.flush()  # ensure DELETE is sent before INSERT
        logger.info(f"Deleted existing marks for course_id={course_id}, inserting {len(students)} new records")

        records = []
        for s in students:
            rec = COAttainment(
                course_id=course_id,
                student_id=s["student_id"],
                student_name=s["student_name"],
            )
            rec.marks = s["marks"]
            self.db.add(rec)
            records.append(rec)

        await self.db.commit()
        logger.info(f"Saved marks for {len(records)} students, course_id={course_id}")

        # Verify save by re-reading first record
        try:
            verify = await self.db.execute(
                select(COAttainment).where(COAttainment.course_id == course_id).limit(1)
            )
            first = verify.scalar_one_or_none()
            if first:
                logger.info(f"Verify OK — first record marks keys: {list(first.marks.keys())}")
            else:
                logger.error(f"Verify FAILED — no records found after save for course_id={course_id}")
        except Exception as ve:
            logger.error(f"Verify check error: {ve}")

        return {
            "course_id": course_id,
            "students_saved": len(records),
            "message": "Marks saved successfully. Call /attainment/report/{course_id} to generate the report.",
        }

    # ------------------------------------------------------------------
    # 2. Calculate attainment
    # ------------------------------------------------------------------

    # ------------------------------------------------------------------
    # Internal helpers for calculate()
    # ------------------------------------------------------------------

    @staticmethod
    def _detect_marks_format(records) -> str:
        """
        Detect the format of the stored marks dicts.

        Returns one of:
          "co_wise"        — keys are CO IDs: {"CO1": {"Quiz": 8, "UT1": 14}, ...}
          "component_wise" — keys are component names: {"Quiz 1": {"Total": 8.5, "Q1": 2.5}, ...}
          "exam_wise_flat" — keys are component names, values are floats: {"Quiz": 8.5, ...}
        """
        import re as _re
        if not records:
            return "component_wise"
        sample = records[0].marks or {}
        if not sample:
            return "component_wise"
        first_key = next(iter(sample))
        first_val = sample[first_key]
        # CO-wise: key looks like CO1, CO2, …
        if _re.match(r'^CO\d+$', str(first_key), _re.IGNORECASE):
            return "co_wise"
        # Flat exam-wise: value is a plain number (not a dict)
        if not isinstance(first_val, dict):
            return "exam_wise_flat"
        # Otherwise it's component-wise (value is a dict with "Total" etc.)
        return "component_wise"

    @staticmethod
    def _student_total_from_marks(marks: dict, fmt: str) -> float:
        """Return the grand total marks for a student given their marks dict and format."""
        if fmt == "co_wise":
            # Sum all component values across all COs
            total = 0.0
            for co_marks in marks.values():
                if isinstance(co_marks, dict):
                    total += sum(float(v) for v in co_marks.values() if v is not None)
                elif co_marks is not None:
                    total += float(co_marks)
            return total
        elif fmt == "component_wise":
            # Each key is a component; value is a dict with at least "Total"
            total = 0.0
            for comp_marks in marks.values():
                if isinstance(comp_marks, dict):
                    t = comp_marks.get("Total")
                    if t is not None:
                        total += float(t)
                elif comp_marks is not None:
                    total += float(comp_marks)
            return total
        else:  # exam_wise_flat
            return sum(float(v) for v in marks.values() if v is not None)

    @staticmethod
    def _max_marks_from_eval_cfg(eval_cfg: dict) -> float:
        """Sum of all component max marks from the evaluation config."""
        components = eval_cfg.get("components", {})
        if not components:
            return 100.0
        return sum(float(v) for v in components.values())

    # ------------------------------------------------------------------
    # 2. Calculate attainment
    # ------------------------------------------------------------------

    async def calculate(self, course_id: int) -> dict:
        """
        Returns full attainment calculation dict without writing any file.

        Handles three marks storage formats produced by parse_marks_xlsx:

        1. **co_wise**        — {"CO1": {"Quiz": 8, "UT": 14}, "CO2": {...}}
           CO attainment is computed per-CO directly.

        2. **component_wise** — {"Quiz 1": {"Total": 8.5}, "Unit Test 1": {"Total": 14}}
           (produced by the multi-sheet template upload)
           Attainment is computed on total marks across all components, then
           applied uniformly to all COs (since the template doesn't carry CO-per-question
           mapping). The CO-PO matrix still drives PO attainment weights.

        3. **exam_wise_flat** — {"Quiz": 8.5, "Unit Test": 14}
           Same as component_wise but values are floats instead of dicts.
        """
        course_svc = CourseService(self.db)
        course = await course_svc.get_course(course_id)

        result = await self.db.execute(
            select(COAttainment).where(COAttainment.course_id == course_id)
        )
        records = result.scalars().all()

        if not records:
            raise OBEException(
                f"No student marks found for course_id={course_id}. Upload marks first.",
                status_code=400,
            )

        cos = course.cos                         # [{co_id, statement, bloom_level}]
        eval_cfg = course.evaluation_config      # {components: {name: max_marks}, ...}
        co_po_matrix = course.co_po_matrix
        pos = course.pos

        components = eval_cfg.get("components", {})  # {"Quiz": 10, "Unit Test": 20, ...}
        all_co_ids = [co["co_id"] for co in cos]

        fmt = self._detect_marks_format(records)
        logger.info(f"Marks format detected for course_id={course_id}: {fmt}")

        threshold = ATTAINMENT_THRESHOLD_PCT / 100

        co_attainment: dict[str, dict] = {}

        if fmt == "co_wise":
            # ── Original CO-wise path ──────────────────────────────────────
            co_max: dict[str, float] = {}
            for co_id in all_co_ids:
                component_names = set()
                for rec in records:
                    co_marks = rec.marks.get(co_id, {})
                    if isinstance(co_marks, dict):
                        component_names.update(co_marks.keys())
                co_max[co_id] = (
                    sum(components.get(c, 10) for c in component_names)
                    or sum(components.values())
                    or 40.0
                )

            co_student_totals: dict[str, list[float]] = defaultdict(list)
            for rec in records:
                for co_id in all_co_ids:
                    co_marks = rec.marks.get(co_id, {})
                    if isinstance(co_marks, dict):
                        total = sum(float(v) for v in co_marks.values() if v is not None)
                    else:
                        total = float(co_marks) if co_marks is not None else 0.0
                    co_student_totals[co_id].append(total)

            for co_id in all_co_ids:
                max_m = co_max.get(co_id, 1)
                totals = co_student_totals[co_id]
                passed = sum(1 for t in totals if t >= threshold * max_m)
                attainment_pct = round((passed / len(totals)) * 100, 2) if totals else 0.0
                avg_marks = round(sum(totals) / len(totals), 2) if totals else 0.0
                co_attainment[co_id] = {
                    "co_id": co_id,
                    "statement": next((c["statement"] for c in cos if c["co_id"] == co_id), ""),
                    "bloom_level": next((c["bloom_level"] for c in cos if c["co_id"] == co_id), ""),
                    "max_marks": max_m,
                    "avg_marks_scored": avg_marks,
                    "students_passed_threshold": passed,
                    "total_students": len(totals),
                    "attainment_percentage": attainment_pct,
                    "attainment_level": self._attainment_level(attainment_pct),
                    "target_met": attainment_pct >= 60,
                }

        else:
            # ── Component-wise or exam-wise flat path ─────────────────────
            # Compute per-student grand total across all components/exams.
            # Determine max possible marks from eval_cfg or from the stored data.
            stored_component_maxes: dict[str, float] = {}
            for rec in records:
                for comp, val in rec.marks.items():
                    if isinstance(val, dict):
                        t = val.get("Total")
                        if t is not None:
                            stored_component_maxes[comp] = max(
                                stored_component_maxes.get(comp, 0), float(t)
                            )
                    elif val is not None:
                        stored_component_maxes[comp] = max(
                            stored_component_maxes.get(comp, 0), float(val)
                        )

            # Max total = sum of eval_cfg component maxes if available, else inferred
            eval_max_total = sum(float(v) for v in components.values()) if components else 0.0
            if eval_max_total <= 0:
                # Fallback: sum of observed max per component × 1.1 (generous ceiling)
                eval_max_total = sum(stored_component_maxes.values()) * 1.1 or 100.0

            # Per-student totals
            student_totals: list[float] = [
                self._student_total_from_marks(rec.marks, fmt) for rec in records
            ]

            passed_global = sum(1 for t in student_totals if t >= threshold * eval_max_total)
            avg_total = round(sum(student_totals) / len(student_totals), 2) if student_totals else 0.0
            global_attainment_pct = (
                round((passed_global / len(student_totals)) * 100, 2) if student_totals else 0.0
            )

            logger.info(
                f"Component-wise attainment: eval_max={eval_max_total}, "
                f"avg={avg_total}, passed={passed_global}/{len(student_totals)}, "
                f"attainment={global_attainment_pct}%"
            )

            # Apply the same attainment % to all COs
            # (without a CO-per-question mapping we cannot split by CO)
            for co_id in all_co_ids:
                co_attainment[co_id] = {
                    "co_id": co_id,
                    "statement": next((c["statement"] for c in cos if c["co_id"] == co_id), ""),
                    "bloom_level": next((c["bloom_level"] for c in cos if c["co_id"] == co_id), ""),
                    "max_marks": round(eval_max_total, 2),
                    "avg_marks_scored": avg_total,
                    "students_passed_threshold": passed_global,
                    "total_students": len(student_totals),
                    "attainment_percentage": global_attainment_pct,
                    "attainment_level": self._attainment_level(global_attainment_pct),
                    "target_met": global_attainment_pct >= 60,
                }

        # ── PO attainment (weighted average via CO-PO matrix) ─────────────
        po_attainment: dict[str, dict] = {}
        for po in pos:
            po_id = po["po_id"]
            weights = []
            for co_id in all_co_ids:
                weight = co_po_matrix.get(co_id, {}).get(po_id, 0)
                if weight > 0:
                    weights.append((co_attainment[co_id]["attainment_percentage"], weight))
            if weights:
                total_w = sum(w for _, w in weights)
                po_pct = round(sum(p * w for p, w in weights) / total_w, 2)
            else:
                po_pct = 0.0
            po_attainment[po_id] = {
                "po_id": po_id,
                "statement": po["statement"],
                "attainment_percentage": po_pct,
                "attainment_level": self._attainment_level(po_pct),
            }

        # ── PSO attainment (same weighted-average formula as PO) ──────────
        psos = course.psos          # [{pso_id, statement}]
        co_pso_matrix = course.co_pso_matrix
        pso_attainment: dict[str, dict] = {}
        for pso in psos:
            pso_id = pso["pso_id"]
            weights = []
            for co_id in all_co_ids:
                weight = co_pso_matrix.get(co_id, {}).get(pso_id, 0)
                if weight > 0:
                    weights.append((co_attainment[co_id]["attainment_percentage"], weight))
            if weights:
                total_w = sum(w for _, w in weights)
                pso_pct = round(sum(p * w for p, w in weights) / total_w, 2)
            else:
                pso_pct = 0.0
            pso_attainment[pso_id] = {
                "pso_id": pso_id,
                "statement": pso["statement"],
                "attainment_percentage": pso_pct,
                "attainment_level": self._attainment_level(pso_pct),
            }

        return {
            "course_id": course_id,
            "course_name": course.course_name,
            "course_code": course.course_code,
            "total_students": len(records),
            "threshold_percentage": ATTAINMENT_THRESHOLD_PCT,
            "marks_format": fmt,
            "co_attainment": co_attainment,
            "po_attainment": po_attainment,
            "pso_attainment": pso_attainment,
            "overall_co_attainment": round(
                sum(v["attainment_percentage"] for v in co_attainment.values()) / len(co_attainment), 2
            ) if co_attainment else 0.0,
        }

    # ------------------------------------------------------------------
    # 3. Generate Word report
    # ------------------------------------------------------------------

    async def generate_report(self, course_id: int) -> dict:
        data = await self.calculate(course_id)

        course_svc = CourseService(self.db)
        course = await course_svc.get_course(course_id)

        _storage = get_storage()
        _filename = f"attainment_report_{course_id}.docx"
        import tempfile as _tmp
        from pathlib import Path as _Path
        with _tmp.TemporaryDirectory() as _t:
            filepath = str(_Path(_t) / _filename)
            self._build_docx(course, data, filepath)
            _storage.save_from_path(_CATEGORY, _filename, _Path(filepath))
        filepath = str(_storage.get_path(_CATEGORY, _filename))

        return {
            "course_id": course_id,
            "course_name": data["course_name"],
            "filename": os.path.basename(filepath),
            "download_url": f"/attainment/download/{course_id}",
            "overall_co_attainment": data["overall_co_attainment"],
            "total_students": data["total_students"],
        }

    # ------------------------------------------------------------------
    # Word document builder
    # ------------------------------------------------------------------

    def _build_docx(self, course: Course, data: dict, filepath: str):
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # ── Title
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("CO & PO ATTAINMENT REPORT")
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = s.add_run(
            f"{data['course_name']}   |   {data['course_code']}   |   "
            f"Students: {data['total_students']}   |   "
            f"Threshold: {data['threshold_percentage']}%"
        )
        sr.font.size = Pt(11)
        sr.bold = True

        doc.add_paragraph()

        # ── CO Attainment Table
        self._heading(doc, "Course Outcome (CO) Attainment")

        co_tbl = doc.add_table(rows=1, cols=7)
        co_tbl.style = "Table Grid"
        for cell, text in zip(
            co_tbl.rows[0].cells,
            ["CO", "Statement", "Bloom's", "Max Marks", "Avg Scored", "Students ≥ Threshold", "Attainment %"],
        ):
            self._hdr_cell(cell, text)

        for co_id, co_data in data["co_attainment"].items():
            row = co_tbl.add_row().cells
            row[0].text = co_id
            row[1].text = co_data["statement"]
            row[2].text = co_data["bloom_level"]
            row[3].text = str(co_data["max_marks"])
            row[4].text = str(co_data["avg_marks_scored"])
            row[5].text = f"{co_data['students_passed_threshold']} / {co_data['total_students']}"
            pct = co_data["attainment_percentage"]
            row[6].text = f"{pct}%  ({co_data['attainment_level']})"
            # Colour attainment cell by level
            fill = "70AD47" if co_data["target_met"] else "FF0000"
            self._shade_cell(row[6], fill)
            row[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for c in row:
                c.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # ── Overall CO Attainment callout
        overall_para = doc.add_paragraph()
        overall_run = overall_para.add_run(
            f"Overall CO Attainment: {data['overall_co_attainment']}%"
        )
        overall_run.bold = True
        overall_run.font.size = Pt(11)
        overall_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        doc.add_paragraph()

        # ── PO Attainment Table
        self._heading(doc, "Program Outcome (PO) Attainment (via CO-PO Matrix)")

        po_tbl = doc.add_table(rows=1, cols=4)
        po_tbl.style = "Table Grid"
        for cell, text in zip(
            po_tbl.rows[0].cells,
            ["PO", "Statement", "Attainment %", "Level"],
        ):
            self._hdr_cell(cell, text, color="7F3F98")

        for po_id, po_data in data["po_attainment"].items():
            row = po_tbl.add_row().cells
            row[0].text = po_id
            row[1].text = po_data["statement"]
            row[2].text = f"{po_data['attainment_percentage']}%"
            row[3].text = po_data["attainment_level"]
            for c in row:
                c.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # ── CO-PO Matrix section
        self._heading(doc, "CO-PO Correlation Matrix (Reference)")
        pos_list = course.pos
        po_ids = [p["po_id"] for p in pos_list]
        co_ids = list(data["co_attainment"].keys())

        matrix_tbl = doc.add_table(rows=1, cols=len(po_ids) + 1)
        matrix_tbl.style = "Table Grid"
        header_cells = matrix_tbl.rows[0].cells
        header_cells[0].text = "CO \\ PO"
        header_cells[0].paragraphs[0].runs[0].bold = True
        for i, po_id in enumerate(po_ids):
            self._hdr_cell(header_cells[i + 1], po_id, color="C55A11")

        for co_id in co_ids:
            row = matrix_tbl.add_row().cells
            r0 = row[0].paragraphs[0].add_run(co_id)
            r0.bold = True
            r0.font.size = Pt(9)
            for i, po_id in enumerate(po_ids):
                val = course.co_po_matrix.get(co_id, {}).get(po_id, 0)
                row[i + 1].text = str(val) if val else "-"
                row[i + 1].paragraphs[0].runs[0].font.size = Pt(9)

        # ── PSO Attainment Table (skip entirely if no PSOs configured)
        if data.get("pso_attainment"):
            doc.add_paragraph()
            self._heading(doc, "Program Specific Outcome (PSO) Attainment (via CO-PSO Matrix)")

            pso_tbl = doc.add_table(rows=1, cols=4)
            pso_tbl.style = "Table Grid"
            for cell, text in zip(
                pso_tbl.rows[0].cells,
                ["PSO", "Statement", "Attainment %", "Level"],
            ):
                self._hdr_cell(cell, text, color="1F6B3A")

            for pso_id, pso_data in data["pso_attainment"].items():
                row = pso_tbl.add_row().cells
                row[0].text = pso_id
                row[1].text = pso_data["statement"]
                row[2].text = f"{pso_data['attainment_percentage']}%"
                row[3].text = pso_data["attainment_level"]
                for c in row:
                    c.paragraphs[0].runs[0].font.size = Pt(9)

            doc.add_paragraph()
            self._heading(doc, "CO-PSO Correlation Matrix (Reference)")
            psos_list = course.psos
            pso_ids = [p["pso_id"] for p in psos_list]
            co_ids_pso = list(data["co_attainment"].keys())

            cpso_tbl = doc.add_table(rows=1, cols=len(pso_ids) + 1)
            cpso_tbl.style = "Table Grid"
            hdr = cpso_tbl.rows[0].cells
            hdr[0].text = "CO \\ PSO"
            hdr[0].paragraphs[0].runs[0].bold = True
            for i, pso_id in enumerate(pso_ids):
                self._hdr_cell(hdr[i + 1], pso_id, color="1F6B3A")

            for co_id in co_ids_pso:
                row = cpso_tbl.add_row().cells
                r0 = row[0].paragraphs[0].add_run(co_id)
                r0.bold = True
                r0.font.size = Pt(9)
                for i, pso_id in enumerate(pso_ids):
                    val = course.co_pso_matrix.get(co_id, {}).get(pso_id, 0)
                    row[i + 1].text = str(val) if val else "-"
                    row[i + 1].paragraphs[0].runs[0].font.size = Pt(9)

        doc.save(filepath)
        logger.info(f"Attainment report saved → {filepath}")

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _attainment_level(pct: float) -> str:
        if pct >= 70:
            return "High"
        if pct >= 50:
            return "Medium"
        return "Low"

    @staticmethod
    def _heading(doc: Document, text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    @staticmethod
    def _hdr_cell(cell, text: str, color: str = "1F497D"):
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        AttainmentService._shade_cell(cell, color)

    @staticmethod
    def _shade_cell(cell, color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    # ------------------------------------------------------------------
    # Template-based report generation (additive — existing generate_report untouched)
    # ------------------------------------------------------------------

    async def generate_report_from_template(self, course_id: int, template_bytes: bytes) -> dict:
        """
        Generate a CO/PO attainment Word report whose section structure follows
        the user-uploaded .docx template.

        Steps:
          1. Extract headings from the template using python-docx
          2. Build the attainment data (same calculate() as normal flow)
          3. Ask Gemini to generate content for each heading section
          4. Write a new .docx with template headings + AI-generated content + data tables
          5. Save via storage (same path as normal report — overrides it)
        """
        import io as _io
        from docx import Document as _Document
        from docx.shared import Pt as _Pt, RGBColor as _RGB
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
        import google.generativeai as genai
        import json as _json

        logger.info(f"Template-based attainment report for course_id={course_id}")

        # ── 1. Extract headings from template ──────────────────────────
        tpl_doc = _Document(_io.BytesIO(template_bytes))
        headings = []
        for para in tpl_doc.paragraphs:
            style = para.style.name.lower()
            text = para.text.strip()
            if text and ('heading' in style or para.runs and any(r.bold for r in para.runs)):
                headings.append(text)

        # Fallback: if no headings detected, use standard NBA sections
        if not headings:
            headings = [
                "Course Information",
                "Course Outcome (CO) Attainment",
                "Program Outcome (PO) Attainment",
                "CO-PO Correlation Matrix",
                "Gap Analysis",
                "Recommendations",
            ]
        logger.info(f"Template headings extracted: {headings}")

        # ── 2. Calculate attainment data ───────────────────────────────
        data = await self.calculate(course_id)
        course_svc = CourseService(self.db)
        course = await course_svc.get_course(course_id)

        # ── 3. Ask Gemini for section content ──────────────────────────
        attainment_summary = _json.dumps({
            "course_name": data["course_name"],
            "course_code": data["course_code"],
            "total_students": data["total_students"],
            "overall_co_attainment": data["overall_co_attainment"],
            "threshold_percentage": data["threshold_percentage"],
            "co_attainment": {
                co_id: {
                    "statement": v["statement"],
                    "attainment_percentage": v["attainment_percentage"],
                    "attainment_level": v["attainment_level"],
                    "target_met": v["target_met"],
                }
                for co_id, v in data["co_attainment"].items()
            },
            "po_attainment": {
                po_id: {
                    "statement": v["statement"],
                    "attainment_percentage": v["attainment_percentage"],
                    "attainment_level": v["attainment_level"],
                }
                for po_id, v in data["po_attainment"].items()
            },
        }, indent=2)

        prompt = f"""You are an NBA/NAAC accreditation report writer for engineering colleges in India.
Generate report content for each section heading listed below, based on the attainment data provided.
For each heading, write 2-4 professional paragraphs suitable for an accreditation report.
Return ONLY a JSON object where keys are the exact heading strings and values are the text content for that section.
Do not include markdown fences.

Headings to fill:
{_json.dumps(headings)}

Attainment data:
{attainment_summary}

Return format (JSON only, no fences):
{{"<heading1>": "<content>", "<heading2>": "<content>", ...}}
"""
        section_content = {}
        try:
            from backend.core.config import GEMINI_API_KEY
            genai.configure(api_key=GEMINI_API_KEY)
            model = genai.GenerativeModel("gemini-1.5-flash")
            resp = model.generate_content(prompt)
            raw = resp.text.strip()
            raw = raw.replace("```json", "").replace("```", "").strip()
            section_content = _json.loads(raw)
        except Exception as e:
            logger.warning(f"Gemini section generation failed: {e} — using fallback summaries")
            for h in headings:
                section_content[h] = (
                    f"This section covers {h.lower()} for {data['course_name']} "
                    f"({data['course_code']}). Overall CO attainment: {data['overall_co_attainment']}%."
                )

        # ── 4. Build the output .docx ──────────────────────────────────
        out_doc = _Document()
        for sec in out_doc.sections:
            from docx.shared import Inches as _Inches
            sec.top_margin = _Inches(1)
            sec.bottom_margin = _Inches(1)
            sec.left_margin = _Inches(1)
            sec.right_margin = _Inches(1)

        # Title
        title_p = out_doc.add_paragraph()
        title_p.alignment = _ALIGN.CENTER
        title_r = title_p.add_run("CO & PO ATTAINMENT REPORT")
        title_r.bold = True
        title_r.font.size = _Pt(18)
        title_r.font.color.rgb = _RGB(0x1F, 0x49, 0x7D)

        sub_p = out_doc.add_paragraph()
        sub_p.alignment = _ALIGN.CENTER
        sub_r = sub_p.add_run(
            f"{data['course_name']}  |  {data['course_code']}  |  "
            f"Students: {data['total_students']}  |  Threshold: {data['threshold_percentage']}%"
        )
        sub_r.font.size = _Pt(11)
        sub_r.bold = True
        out_doc.add_paragraph()

        # Sections from template headings
        for heading in headings:
            # Heading
            h_para = out_doc.add_paragraph()
            h_run = h_para.add_run(heading)
            h_run.bold = True
            h_run.font.size = _Pt(12)
            h_run.font.color.rgb = _RGB(0x1F, 0x49, 0x7D)
            h_para.paragraph_format.space_before = _Pt(10)
            h_para.paragraph_format.space_after = _Pt(4)

            # AI-generated content
            content = section_content.get(heading, "")
            if content:
                c_para = out_doc.add_paragraph()
                c_run = c_para.add_run(content)
                c_run.font.size = _Pt(10)
                out_doc.add_paragraph()

            # Inject data tables for known section types
            heading_lower = heading.lower()
            if any(kw in heading_lower for kw in ["co attainment", "course outcome"]):
                self._build_docx_co_table(out_doc, data)
                out_doc.add_paragraph()
            elif any(kw in heading_lower for kw in ["pso attainment", "program specific"]):
                self._build_docx_pso_table(out_doc, data)
                out_doc.add_paragraph()
            elif any(kw in heading_lower for kw in ["po attainment", "program outcome"]):
                self._build_docx_po_table(out_doc, data)
                out_doc.add_paragraph()
            elif any(kw in heading_lower for kw in ["matrix", "correlation"]):
                self._build_docx_matrix_table(out_doc, course, data)
                out_doc.add_paragraph()

        # ── 5. Save via storage ────────────────────────────────────────
        _storage = get_storage()
        _filename = f"attainment_report_{course_id}.docx"
        import tempfile as _tmp
        from pathlib import Path as _Path
        with _tmp.TemporaryDirectory() as _t:
            fp = str(_Path(_t) / _filename)
            out_doc.save(fp)
            _storage.save_from_path(_CATEGORY, _filename, _Path(fp))
        filepath = str(_storage.get_path(_CATEGORY, _filename))
        logger.info(f"Template-based attainment report saved → {filepath}")

        return {
            "course_id": course_id,
            "course_name": data["course_name"],
            "filename": _filename,
            "download_url": f"/attainment/download/{course_id}",
            "overall_co_attainment": data["overall_co_attainment"],
            "total_students": data["total_students"],
            "template_used": True,
            "sections_generated": headings,
        }

    def _build_docx_co_table(self, doc, data: dict):
        """Shared helper: CO attainment table — reused by template flow."""
        from docx.shared import Pt as _Pt, RGBColor as _RGB
        co_tbl = doc.add_table(rows=1, cols=5)
        co_tbl.style = "Table Grid"
        for cell, text in zip(
            co_tbl.rows[0].cells,
            ["CO", "Statement", "Max Marks", "Attainment %", "Level"],
        ):
            self._hdr_cell(cell, text)
        for co_id, co_data in data["co_attainment"].items():
            row = co_tbl.add_row().cells
            row[0].text = co_id
            row[1].text = co_data["statement"]
            row[2].text = str(co_data["max_marks"])
            pct = co_data["attainment_percentage"]
            row[3].text = f"{pct}%"
            row[4].text = co_data["attainment_level"]
            fill = "70AD47" if co_data["target_met"] else "FF6B6B"
            self._shade_cell(row[3], fill)
            row[3].paragraphs[0].runs[0].font.color.rgb = _RGB(0xFF, 0xFF, 0xFF)
            for c in row:
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = _Pt(9)

    def _build_docx_po_table(self, doc, data: dict):
        """Shared helper: PO attainment table — reused by template flow."""
        from docx.shared import Pt as _Pt
        po_tbl = doc.add_table(rows=1, cols=4)
        po_tbl.style = "Table Grid"
        for cell, text in zip(
            po_tbl.rows[0].cells,
            ["PO", "Statement", "Attainment %", "Level"],
        ):
            self._hdr_cell(cell, text, color="7F3F98")
        for po_id, po_data in data["po_attainment"].items():
            row = po_tbl.add_row().cells
            row[0].text = po_id
            row[1].text = po_data["statement"]
            row[2].text = f"{po_data['attainment_percentage']}%"
            row[3].text = po_data["attainment_level"]
            for c in row:
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = _Pt(9)

    def _build_docx_pso_table(self, doc, data: dict):
        """Shared helper: PSO attainment table — green header to distinguish from PO."""
        from docx.shared import Pt as _Pt
        if not data.get("pso_attainment"):
            p = doc.add_paragraph()
            p.add_run("No PSOs configured for this course.").italic = True
            return
        pso_tbl = doc.add_table(rows=1, cols=4)
        pso_tbl.style = "Table Grid"
        for cell, text in zip(
            pso_tbl.rows[0].cells,
            ["PSO", "Statement", "Attainment %", "Level"],
        ):
            self._hdr_cell(cell, text, color="1F6B3A")
        for pso_id, pso_data in data["pso_attainment"].items():
            row = pso_tbl.add_row().cells
            row[0].text = pso_id
            row[1].text = pso_data["statement"]
            row[2].text = f"{pso_data['attainment_percentage']}%"
            row[3].text = pso_data["attainment_level"]
            for c in row:
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = _Pt(9)

    def _build_docx_matrix_table(self, doc, course, data: dict):
        """Shared helper: CO-PO matrix table — reused by template flow."""
        from docx.shared import Pt as _Pt
        pos_list = course.pos
        po_ids = [p["po_id"] for p in pos_list]
        co_ids = list(data["co_attainment"].keys())
        if not po_ids or not co_ids:
            return
        matrix_tbl = doc.add_table(rows=1, cols=len(po_ids) + 1)
        matrix_tbl.style = "Table Grid"
        header_cells = matrix_tbl.rows[0].cells
        header_cells[0].text = "CO \\ PO"
        if header_cells[0].paragraphs[0].runs:
            header_cells[0].paragraphs[0].runs[0].bold = True
        for i, po_id in enumerate(po_ids):
            self._hdr_cell(header_cells[i + 1], po_id, color="C55A11")
        for co_id in co_ids:
            row = matrix_tbl.add_row().cells
            r0 = row[0].paragraphs[0].add_run(co_id)
            r0.bold = True
            r0.font.size = _Pt(9)
            for i, po_id in enumerate(po_ids):
                val = course.co_po_matrix.get(co_id, {}).get(po_id, 0)
                row[i + 1].text = str(val) if val else "-"
                if row[i + 1].paragraphs[0].runs:
                    row[i + 1].paragraphs[0].runs[0].font.size = _Pt(9)

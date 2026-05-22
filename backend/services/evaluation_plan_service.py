# backend/services/evaluation_plan_service.py
"""
Generates evaluation plan matching the exact SIT format:
Table: Sr.No | Component | Unit Syllabus | CO | Marks | Weightage | Tentative Date
"""
import json, os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.ext.asyncio import AsyncSession
from backend.core.exceptions import LLMError
from backend.core.llm import get_llm_response
from backend.core.logger import get_logger
from backend.services.course_service import CourseService

logger = get_logger(__name__)
from backend.core.storage import get_storage
_CATEGORY = "evaluation_plans"

_NAVY  = "1F3864"
_LIGHT = "D6DCE4"
_WHITE = RGBColor(0xFF,0xFF,0xFF)
_BLACK = RGBColor(0,0,0)
_GREEN = RGBColor(0x00,0x80,0x00)
_FONT  = "Times New Roman"

class EvaluationPlanService:
    def __init__(self, db: AsyncSession):
        self.db = db

    @staticmethod
    def get_filepath(course_id: int) -> str:
        storage = get_storage()
        p = storage.get_path(_CATEGORY, f"evaluation_plan_{course_id}.docx")
        return str(p) if p else str(get_storage()._dir(_CATEGORY) / f"evaluation_plan_{course_id}.docx")

    async def generate(self, course_id: int) -> dict:
        course_svc = CourseService(self.db)
        course = await course_svc.get_course(course_id)
        course_name   = course.course_name
        course_code   = course.course_code
        faculty_name  = course.faculty_name
        department    = course.department
        semester      = course.semester
        academic_year = course.academic_year
        cos           = course.cos
        eval_cfg      = course.evaluation_config
        eval_cfg      = {**eval_cfg, "credits": str(course.credits)}
        logger.info(f"Generating evaluation plan for '{course_name}' ({course_code})")
        prompt = self._build_prompt(course_name, course_code, cos, eval_cfg)
        plan   = await self._call_llm(prompt)
        _storage  = get_storage()
        _filename = f"evaluation_plan_{course_id}.docx"
        filepath  = self._build_docx(
            course_name, course_code, faculty_name, department,
            semester, academic_year, cos, eval_cfg, plan, _storage, _filename
        )
        return {
            "course_id":       course_id,
            "course_name":     course_name,
            "filename":        os.path.basename(filepath),
            "download_url":    f"/evaluation-plan/download/{course_id}",
            "cie_total":       eval_cfg.get("continuous_assessment_total", 30),
            "see_total":       eval_cfg.get("end_sem_total", 60),
            "evaluation_plan": plan,
        }

    def _build_prompt(self, course_name, course_code, cos, eval_cfg):
        cos_text = "\n".join(f"  {c['co_id']}: {c['statement']}" for c in cos)
        co_ids   = [c["co_id"] for c in cos]
        components = eval_cfg.get("components", {})
        comp_text  = "\n".join(f"  - {k}: {v} marks" for k,v in components.items())
        cie_total  = eval_cfg.get("continuous_assessment_total", 30)
        num_units  = len(cos)

        return f"""You are an OBE evaluation expert for engineering colleges.

Course: {course_name} ({course_code})
Course Outcomes: {co_ids}
{cos_text}

CIE Total: {cie_total} marks
Components:
{comp_text}

Generate an EVALUATION PLAN with exactly these CA components (CA1, CA2, CA3...) mapped to units and COs.
Each component must specify: which units it covers, which COs, marks, weightage %, and a tentative date (month 2026).

Return ONLY valid JSON, no markdown:
{{"ca_components":[{{"sr_no":"CA1","component":"Experiential Learning","unit_syllabus":"Unit 1, Unit 2","co_mapped":"CO1, CO2","marks":10,"weightage":"33%","tentative_date":"February 2026"}}]}}"""

    async def _call_llm(self, prompt: str) -> dict:
        raw = await get_llm_response(prompt)
        if raw.startswith("```"):
            raw = "\n".join(raw.split("\n")[1:-1])
        raw = raw.strip()
        try:
            return json.loads(raw)
        except json.JSONDecodeError as e:
            logger.error(f"LLM JSON parse error: {e}")
            raise LLMError("LLM returned invalid JSON for evaluation plan")

    def _build_docx(self, course_name, course_code, faculty_name, department,
                    semester, academic_year, cos, eval_cfg, data, _storage, _filename) -> str:
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(0.5)
            sec.left_margin = sec.right_margin = Inches(0.6)

        cia_total = eval_cfg.get("continuous_assessment_total", 30)
        ese_total = eval_cfg.get("end_sem_total", 45)
        credits   = eval_cfg.get("credits", "")

        def _para(text="", bold=False, size=12, color=None, align=None, indent=None):
            p = doc.add_paragraph()
            if align:
                p.alignment = align
            if indent:
                p.paragraph_format.left_indent = Inches(indent)
            if text:
                r = p.add_run(text)
                r.bold = bold
                r.font.name = _FONT
                r.font.size = Pt(size)
                r.font.color.rgb = color if color else _BLACK
            return p

        def _split_para(left_bold_text, left_normal_text, right_bold_text, right_normal_text, size=10):
            p = doc.add_paragraph()
            if left_bold_text:
                rb = p.add_run(left_bold_text)
                rb.bold = True; rb.font.name = _FONT; rb.font.size = Pt(size); rb.font.color.rgb = _BLACK
            if left_normal_text:
                rn = p.add_run(left_normal_text)
                rn.font.name = _FONT; rn.font.size = Pt(size); rn.font.color.rgb = _BLACK
            p.add_run("\t")
            if right_bold_text:
                rb2 = p.add_run(right_bold_text)
                rb2.bold = True; rb2.font.name = _FONT; rb2.font.size = Pt(size); rb2.font.color.rgb = _BLACK
            if right_normal_text:
                rn2 = p.add_run(right_normal_text)
                rn2.font.name = _FONT; rn2.font.size = Pt(size); rn2.font.color.rgb = _BLACK
            return p

        # ── FORMAT LABEL (top-right) ──────────────────────────────────
        fmt_p = doc.add_paragraph()
        fmt_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fmt_r = fmt_p.add_run("Format 6")
        fmt_r.font.name = _FONT
        fmt_r.font.size = Pt(12)
        fmt_r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # ── INSTITUTE NAME ────────────────────────────────────────────
        inst_p = doc.add_paragraph()
        inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inst_r = inst_p.add_run("Symbiosis Institute of Technology, Pune")
        inst_r.bold = True; inst_r.font.name = _FONT; inst_r.font.size = Pt(13); inst_r.font.color.rgb = _BLACK

        # ── TITLE ─────────────────────────────────────────────────────
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_r = title_p.add_run("Evaluation Plan")
        title_r.bold = True; title_r.font.name = _FONT; title_r.font.size = Pt(12); title_r.font.color.rgb = _BLACK

        doc.add_paragraph()

        # ── DEPT | BATCH ──────────────────────────────────────────────
        _split_para("\tDepartment: ", department + "  " * 38, "Batch: ", academic_year)

        doc.add_paragraph()

        # ── COURSE NAME | CREDIT ──────────────────────────────────────
        _split_para("\tCourse name:", f" {course_name}", "Credit: ", credits)

        # ── YEAR | SEM ────────────────────────────────────────────────
        _split_para("\tYear:", f" {academic_year}", "Sem: ", semester)

        # ── FACULTY ───────────────────────────────────────────────────
        fac_p = doc.add_paragraph()
        fb = fac_p.add_run("Name of the faculty member: ")
        fb.bold = True; fb.font.name = _FONT; fb.font.size = Pt(10); fb.font.color.rgb = _BLACK
        fn = fac_p.add_run(faculty_name)
        fn.font.name = _FONT; fn.font.size = Pt(10); fn.font.color.rgb = _BLACK

        doc.add_paragraph()

        # ── CA / ESE MARKS ────────────────────────────────────────────
        ca_p = doc.add_paragraph()
        ca_bold = ca_p.add_run("CA")
        ca_bold.bold = True; ca_bold.font.name = _FONT; ca_bold.font.size = Pt(10); ca_bold.font.color.rgb = _BLACK
        ca_rest = ca_p.add_run(f" -{cia_total} marks")
        ca_rest.font.name = _FONT; ca_rest.font.size = Pt(10); ca_rest.font.color.rgb = _BLACK

        ese_p = doc.add_paragraph()
        ese_r = ese_p.add_run(f"ESE – {ese_total} marks")
        ese_r.font.name = _FONT; ese_r.font.size = Pt(10); ese_r.font.color.rgb = _BLACK

        doc.add_paragraph()

        # ── COURSE OUTCOMES LIST ──────────────────────────────────────
        co_title = doc.add_paragraph()
        ct = co_title.add_run("Course Outcomes-")
        ct.bold = True; ct.font.name = _FONT; ct.font.size = Pt(10); ct.font.color.rgb = _BLACK

        for co in cos:
            co_p = doc.add_paragraph(style="List Bullet")
            co_r = co_p.add_run(co.get("statement", ""))
            co_r.font.name = _FONT; co_r.font.size = Pt(10); co_r.font.color.rgb = _BLACK

        doc.add_paragraph()

        # ── SECTION TITLE ─────────────────────────────────────────────
        theory_p = doc.add_paragraph()
        tt = theory_p.add_run("Theory Evaluation Components")
        tt.bold = True; tt.font.name = _FONT; tt.font.size = Pt(10); tt.font.color.rgb = _BLACK

        doc.add_paragraph()

        # ── EVALUATION TABLE ──────────────────────────────────────────
        col_headers = ["Sr. No.", "Component", "Unit Syllabus", "CO", "Marks", "Weightage", "Tentative Date"]
        col_widths  = [Inches(0.5), Inches(1.5), Inches(1.8), Inches(1.0), Inches(0.55), Inches(0.75), Inches(1.4)]

        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = "Table Grid"
        for i, (cell, text) in enumerate(zip(tbl.rows[0].cells, col_headers)):
            self._shade(cell, _NAVY)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text); r.bold = True; r.font.name = _FONT; r.font.size = Pt(9); r.font.color.rgb = _WHITE
            cell.width = col_widths[i]

        components = data.get("ca_components", [])
        for i, comp in enumerate(components):
            row = tbl.add_row().cells
            if i % 2 == 1:
                for c in row: self._shade(c, _LIGHT)
            values = [
                comp.get("sr_no", ""),
                comp.get("component", ""),
                comp.get("unit_syllabus", ""),
                comp.get("co_mapped", ""),
                str(comp.get("marks", "")),
                comp.get("weightage", ""),
                comp.get("tentative_date", ""),
            ]
            for j, (c, v) in enumerate(zip(row, values)):
                p = c.paragraphs[0]; p.clear()
                r = p.add_run(v); r.font.name = _FONT; r.font.size = Pt(9); r.font.color.rgb = _BLACK
                if j in (0, 4, 5): p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.width = col_widths[j]

        doc.add_paragraph()

        # ── CA DETAIL NOTES ───────────────────────────────────────────
        for comp in components:
            det_p = doc.add_paragraph()
            label_r = det_p.add_run(f"- ")
            label_r.font.name = _FONT; label_r.font.size = Pt(10); label_r.font.color.rgb = _BLACK
            sr = det_p.add_run(f"{comp.get('sr_no','')}")
            sr.bold = True; sr.font.name = _FONT; sr.font.size = Pt(10); sr.font.color.rgb = _BLACK
            colon = det_p.add_run(":")
            colon.bold = True; colon.font.name = _FONT; colon.font.size = Pt(10); colon.font.color.rgb = _BLACK
            comp_name = det_p.add_run(f" {comp.get('component','')} - {comp.get('marks','')} Marks ({comp.get('co_mapped','')})")
            comp_name.bold = True; comp_name.font.name = _FONT; comp_name.font.size = Pt(10); comp_name.font.color.rgb = _BLACK

            det2 = doc.add_paragraph()
            det2_r = det2.add_run(f"{comp.get('unit_syllabus','')}")
            det2_r.font.name = _FONT; det2_r.font.size = Pt(10); det2_r.font.color.rgb = _BLACK

        doc.add_paragraph()

        # ── SIGNATURE ROW ─────────────────────────────────────────────
        sign_p = doc.add_paragraph()
        sign_r = sign_p.add_run(f"Sign of the faculty member: {faculty_name}")
        sign_r.font.name = _FONT; sign_r.font.size = Pt(10); sign_r.font.color.rgb = _BLACK

        doc.add_paragraph()

        hod_heading_p = doc.add_paragraph()
        hod_heading_r = hod_heading_p.add_run("Sign of HoD: ")
        hod_heading_r.bold = True
        hod_heading_r.font.name = _FONT; hod_heading_r.font.size = Pt(10); hod_heading_r.font.color.rgb = _BLACK

        import tempfile as _tmp
        with _tmp.TemporaryDirectory() as _t:
            _p = Path(_t) / _filename
            doc.save(str(_p))
            _storage.save_from_path(_CATEGORY, _filename, _p)
        filepath = str(_storage.get_path(_CATEGORY, _filename))
        logger.info(f"Evaluation plan saved -> {filepath}")
        return filepath

    @staticmethod
    def _shade(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
        tcPr.append(shd)

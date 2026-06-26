# backend/services/session_plan_service.py
"""
Generates session plan matching the exact SIT (Symbiosis Institute of Technology)
branded format:
  Header: Institute | Dept | Session Plan | Course info rows
  Table columns: Lect.No | Unit No. | Points to Cover | Methodology |
                 Faculty Conducting | Lecture/Exp.Learning/Evaluation | CO
"""
import json, os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from backend.core.exceptions import LLMError
from backend.core.llm import get_llm_response
from backend.core.logger import get_logger
from backend.database.models import EvalPlanRow
from backend.services.course_service import CourseService

logger = get_logger(__name__)
from backend.core.storage import get_storage
_CATEGORY = "session_plans"

_NAVY  = "1F3864"
_LIGHT = "D6DCE4"
_WHITE = RGBColor(0xFF,0xFF,0xFF)
_NAVY_R = RGBColor(0x1F,0x38,0x64)
_BLACK = RGBColor(0,0,0)
_FONT = "Times New Roman"

class SessionPlanService:
    def __init__(self, db: AsyncSession):
        self.db = db

    @staticmethod
    def get_filepath(course_id: int) -> str:
        storage = get_storage()
        p = storage.get_path(_CATEGORY, f"session_plan_{course_id}.docx")
        return str(p) if p else str(get_storage()._dir(_CATEGORY) / f"session_plan_{course_id}.docx")

    async def generate(self, course_id: int) -> dict:
        course_svc = CourseService(self.db)
        course = await course_svc.get_course(course_id)
        course_name   = course.course_name
        course_code   = course.course_code
        faculty_name  = course.faculty_name
        department    = course.department
        semester      = course.semester
        academic_year = course.academic_year
        credits       = course.credits
        cos           = course.cos
        logger.info(f"Generating session plan for '{course_name}' ({course_code})")
        # Total lectures = credits × 15 (standard Indian engineering norm)
        total_lectures = int(credits) * 15

        # Pull eval plan components from DB so session plan places them correctly
        eval_components = []
        try:
            ep_result = await self.db.execute(
                select(EvalPlanRow).where(EvalPlanRow.course_id == course_id)
            )
            ep_row = ep_result.scalar_one_or_none()
            if ep_row and ep_row.rows:
                eval_components = ep_row.rows  # list of dicts with sr, component, unit_syllabus, co, marks, weightage, date
        except Exception as e:
            logger.warning(f"Could not fetch eval plan for session plan generation: {e}")

        prompt = self._build_prompt(course_name, course_code, cos, total_lectures, eval_components)
        plan   = await self._call_llm(prompt)
        _storage  = get_storage()
        _filename = f"session_plan_{course_id}.docx"
        filepath  = self._build_docx(
            course_name, course_code, faculty_name, department,
            semester, academic_year, credits, cos, plan, _storage, _filename
        )
        total_sessions = sum(len(u.get("sessions",[])) for u in plan.get("units",[]))
        return {
            "course_id":      course_id,
            "course_name":    course_name,
            "filename":       os.path.basename(filepath),
            "download_url":   f"/session-plan/download/{course_id}",
            "total_sessions": total_sessions,
            "units":          plan.get("units",[]),
        }

    def _build_prompt(self, course_name, course_code, cos, total_lectures, eval_components=None):
        cos_text = "\n".join(f"  {c['co_id']}: {c['statement']} [Bloom: {c['bloom_level']}]" for c in cos)
        co_ids   = [c["co_id"] for c in cos]
        num_units = len(cos)
        lectures_per_unit = total_lectures // num_units if num_units else total_lectures // 5
        # Build the assessments block from the saved eval plan
        if eval_components:
            assessments_lines = []
            for ec in eval_components:
                sr   = ec.get('sr', ec.get('sr_no', ''))
                comp = ec.get('component', ec.get('comp', ''))
                units_cov = ec.get('unit_syllabus', ec.get('units', ''))
                co_m  = ec.get('co', ec.get('co_mapped', ''))
                marks = ec.get('marks', '')
                date  = ec.get('date', ec.get('tentative_date', ''))
                assessments_lines.append(
                    f"  {sr}: {comp} | Units covered: {units_cov} | CO: {co_m} | Marks: {marks} | Date: {date}"
                )
            assessments_block = (
                "\nSaved Evaluation Plan (MUST be embedded in session plan):\n"
                + "\n".join(assessments_lines)
                + "\n- Each assessment above must appear as a session row with type=\"Evaluation\" "
                  "placed WITHIN the unit(s) it covers, at a natural position (not all at end).\n"
                  "- Use the component name exactly as given (e.g. \"Quiz\", \"Unit Test\", \"Assignment\").\n"
                  "- Map the session to the FIRST CO listed in its CO field.\n"
            )
        else:
            assessments_block = (
                "\n- Include at least 1 Quiz or Unit Test row per unit (type=\"Evaluation\").\n"
            )

        return f"""You are an expert curriculum designer for engineering colleges using OBE.

Course: {course_name} ({course_code})
Total lectures required: {total_lectures} (MANDATORY — generate EXACTLY {total_lectures} session rows total across all units)
Course Outcomes:
{cos_text}
{assessments_block}
Generate a complete SESSION PLAN with EXACTLY {total_lectures} sessions numbered 1 to {total_lectures} sequentially.
Rules:
- Split into {num_units} units. Each unit ~{lectures_per_unit} sessions (distribute evenly; last unit absorbs remainder).
- session_number MUST be a CONTINUOUS GLOBAL counter from 1 to {total_lectures} — NEVER reset between units.
- Every session maps to exactly one CO from: {co_ids}
- teaching_method: one of: Classroom Teaching, Tutorial, Flipped Classroom, Case Study, Problem Solving, Group Discussion
- type: one of: Lecture, Exp. Learning, Evaluation
- Include at least 1 "Exp. Learning" row per unit (topic="Experiential Learning on <unit topic>", type="Exp. Learning")
- Return ONLY valid JSON, no markdown, no extra text.

Schema:
{{"units":[{{"unit_number":1,"unit_title":"string","sessions":[{{"session_number":1,"topic":"string","teaching_method":"Classroom Teaching","type":"Lecture","co_mapped":"CO1"}}]}}]}}"""

    async def _call_llm(self, prompt: str) -> dict:
        raw = await get_llm_response(prompt)
        if raw.startswith("```"):
            raw = "\n".join(raw.split("\n")[1:-1])
        raw = raw.strip()
        try:
            return json.loads(raw)
        except json.JSONDecodeError as e:
            logger.error(f"LLM JSON parse error: {e}")
            raise LLMError("LLM returned invalid JSON for session plan")

    def _build_docx(self, course_name, course_code, faculty_name, department,
                    semester, academic_year, credits, cos, data, _storage, _filename) -> str:
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(0.5)
            sec.left_margin = sec.right_margin = Inches(0.6)

        # ── FORMAT LABEL (top-right, outside table) ───────────────────
        fmt_p = doc.add_paragraph()
        fmt_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fmt_r = fmt_p.add_run("Format: 5")
        fmt_r.font.name = _FONT
        fmt_r.font.size = Pt(12)
        fmt_r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # ── HEADER TABLE — matches template exactly ───────────────────
        hdr = doc.add_table(rows=0, cols=9)
        hdr.style = "Table Grid"

        def _set_cell_borders(cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "auto")
                tcBorders.append(b)
            tcPr.append(tcBorders)

        def _full_row(text, size=11, bold=True, center=True):
            row = hdr.add_row()
            c = row.cells[0]
            for other in row.cells[1:]:
                c = c.merge(other)
            _set_cell_borders(c)
            p = c.paragraphs[0]
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.bold = bold
            r.font.name = _FONT
            r.font.size = Pt(size)
            r.font.color.rgb = _BLACK

        # Row 1: Institute name
        _full_row("Symbiosis Institute of Technology, Pune", size=13)
        # Row 2: "Session Plan" title
        _full_row("Session Plan", size=12)
        # Row 3: Department
        _full_row(f"              Name of the Department – {department}", size=10, center=False)
        # Row 4: Course name | Credit
        r4 = hdr.add_row()
        lc4 = r4.cells[0].merge(r4.cells[5])
        rc4 = r4.cells[6].merge(r4.cells[8])
        for c in [lc4, rc4]:
            _set_cell_borders(c)
        lp4 = lc4.paragraphs[0]
        lr4 = lp4.add_run(f"              Name of the course– {course_name}")
        lr4.bold = True; lr4.font.name = _FONT; lr4.font.size = Pt(10); lr4.font.color.rgb = _BLACK
        rp4 = rc4.paragraphs[0]
        rr4 = rp4.add_run(f"Credit - {credits}")
        rr4.bold = True; rr4.font.name = _FONT; rr4.font.size = Pt(10); rr4.font.color.rgb = _BLACK
        # Row 5: Semester | blank | Batch
        r5 = hdr.add_row()
        lc5 = r5.cells[0].merge(r5.cells[4])
        mc5 = r5.cells[5]
        rc5 = r5.cells[6].merge(r5.cells[8])
        for c in [lc5, mc5, rc5]:
            _set_cell_borders(c)
        sem_run = lc5.paragraphs[0].add_run(f"              Semester –")
        sem_run.bold = True; sem_run.font.name = _FONT; sem_run.font.size = Pt(10); sem_run.font.color.rgb = _BLACK
        sem_val = mc5.paragraphs[0].add_run(f"{semester}")
        sem_val.bold = True; sem_val.font.name = _FONT; sem_val.font.size = Pt(10); sem_val.font.color.rgb = _BLACK
        batch_run = rc5.paragraphs[0].add_run(f"Batch – {academic_year}")
        batch_run.bold = True; batch_run.font.name = _FONT; batch_run.font.size = Pt(10); batch_run.font.color.rgb = _BLACK
        # Row 6: Faculty name
        r6 = hdr.add_row()
        fc = r6.cells[0].merge(r6.cells[8])
        _set_cell_borders(fc)
        fac_run = fc.paragraphs[0].add_run(f"              Name of the faculty– {faculty_name}")
        fac_run.bold = True; fac_run.font.name = _FONT; fac_run.font.size = Pt(10); fac_run.font.color.rgb = _BLACK

        doc.add_paragraph()

        # ── SESSION TABLE ─────────────────────────────────────────────
        col_headers = ["Lect.\nNo","Unit\nNo.","Points to Cover","Methodology",
                       "Faculty Conducting","Lecture/Exp.\nLearning/Evaluation","CO"]
        col_widths  = [Inches(0.45),Inches(0.45),Inches(2.7),Inches(1.3),
                       Inches(1.5),Inches(1.05),Inches(0.45)]

        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = "Table Grid"
        for i,(cell,text) in enumerate(zip(tbl.rows[0].cells, col_headers)):
            self._shade(cell, _NAVY)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text); r.bold=True; r.font.name=_FONT; r.font.size=Pt(9); r.font.color.rgb=_WHITE
            cell.width = col_widths[i]

        lect_num = 0
        for unit in data.get("units",[]):
            for s in unit.get("sessions",[]):
                lect_num += 1
                row = tbl.add_row().cells
                stype = s.get("type","Lecture")
                is_eval = stype in ("Evaluation","Exp. Learning")

                if lect_num % 2 == 0:
                    for c in row: self._shade(c, _LIGHT)

                values = [
                    str(lect_num),  # always global sequential, never per-unit reset
                    str(unit.get("unit_number","")) if not is_eval else "",
                    s.get("topic",""),
                    s.get("teaching_method","Classroom Teaching"),
                    faculty_name,
                    stype,
                    s.get("co_mapped",""),
                ]
                for i,(c,v) in enumerate(zip(row,values)):
                    p = c.paragraphs[0]
                    p.clear()
                    r = p.add_run(v); r.font.name=_FONT; r.font.size = Pt(9)
                    if i in (0,1,6): p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if is_eval: r.bold=True; r.font.color.rgb=_NAVY_R
                    c.width = col_widths[i]

        import tempfile as _tmp
        with _tmp.TemporaryDirectory() as _t:
            _p = Path(_t) / _filename
            doc.save(str(_p))
            _storage.save_from_path(_CATEGORY, _filename, _p)
        filepath = str(_storage.get_path(_CATEGORY, _filename))
        logger.info(f"Session plan saved -> {filepath}")
        return filepath

    @staticmethod
    def _shade(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
        tcPr.append(shd)

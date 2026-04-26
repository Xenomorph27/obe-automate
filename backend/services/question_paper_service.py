# backend/services/question_paper_service.py
import tempfile
from pathlib import Path
from datetime import datetime
from backend.core.logger import get_logger
from backend.core.exceptions import OBEException
from backend.core.storage import get_storage

logger = get_logger(__name__)

CATEGORY = "question_papers"


def _docx(course_name, course_code, duration, total_marks, questions, out_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1.2); s.right_margin = Inches(1.2)
    h = doc.add_heading(course_name, 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Code: {course_code}  |  Duration: {duration}h  |  Total Marks: {total_marks}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("─" * 80)
    doc.add_paragraph("Instructions: Answer ALL questions.")
    sections = {}
    for q in questions:
        sections.setdefault(q.get("section", "Section A"), []).append(q)
    for sec, qs in sections.items():
        doc.add_heading(sec, level=1)
        for q in qs:
            p = doc.add_paragraph()
            p.add_run(f"Q{q.get('question_number', '')}. ").bold = True
            p.add_run(q.get("question_text", ""))
            m = p.add_run(f"  [{q.get('marks', 0)}M | {q.get('bloom_label', '')} | {q.get('co_id', '')}]")
            m.font.color.rgb = RGBColor(0x64, 0x74, 0x8b); m.font.size = Pt(9)
            if q.get("options"):
                for opt in q["options"]:
                    doc.add_paragraph(f"    {opt}", style="List Bullet")
            doc.add_paragraph("")
    doc.save(str(out_path))


def _pdf(course_name, course_code, duration, total_marks, questions, out_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER
    doc = SimpleDocTemplate(str(out_path), pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    NAVY = colors.HexColor("#0f1b2d"); AMBER = colors.HexColor("#f59e0b"); SLATE = colors.HexColor("#64748b")
    ts = ParagraphStyle("T", fontSize=16, textColor=NAVY, alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=4)
    ss = ParagraphStyle("S", fontSize=10, textColor=SLATE, alignment=TA_CENTER, spaceAfter=2)
    hs = ParagraphStyle("H", fontSize=12, textColor=NAVY, fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=6)
    qs = ParagraphStyle("Q", fontSize=10, leading=14, spaceBefore=4, spaceAfter=2)
    ms = ParagraphStyle("M", fontSize=8, textColor=SLATE, spaceAfter=6)
    os_ = ParagraphStyle("O", fontSize=9, leftIndent=20, spaceAfter=2)
    story = [
        Paragraph(course_name, ts),
        Paragraph(f"Code:{course_code} | {duration}h | {total_marks}M", ss),
        Paragraph(f"Date:{datetime.now().strftime('%d %B %Y')}", ss),
        HRFlowable(width="100%", thickness=1, color=AMBER, spaceAfter=8),
        Paragraph("Instructions: Answer ALL questions.", ms),
        Spacer(1, 0.3*cm),
    ]
    sections = {}
    for q in questions:
        sections.setdefault(q.get("section", "Section A"), []).append(q)
    for sec, qs_ in sections.items():
        story += [Paragraph(sec, hs), HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey)]
        for q in qs_:
            story.append(Paragraph(f"<b>Q{q.get('question_number', '')}.</b> {q.get('question_text', '')}", qs))
            story.append(Paragraph(f"[{q.get('marks', 0)}M | {q.get('bloom_label', '')} | {q.get('co_id', '')}]", ms))
            if q.get("options"):
                for opt in q["options"]:
                    story.append(Paragraph(opt, os_))
            story.append(Spacer(1, 0.2*cm))
    doc.build(story)


class QuestionPaperService:
    def generate_documents(self, course_id, course_name, course_code, duration, total_marks, questions) -> dict:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_name = f"qpaper_{course_id}_{ts}.docx"
        pdf_name  = f"qpaper_{course_id}_{ts}.pdf"
        storage = get_storage()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_docx = Path(tmp) / docx_name
            tmp_pdf  = Path(tmp) / pdf_name
            try:
                _docx(course_name, course_code, duration, total_marks, questions, tmp_docx)
            except Exception as e:
                raise OBEException(f"Word generation failed: {e}", 500)
            try:
                _pdf(course_name, course_code, duration, total_marks, questions, tmp_pdf)
            except Exception as e:
                raise OBEException(f"PDF generation failed: {e}", 500)

            docx_path = storage.save_from_path(CATEGORY, docx_name, tmp_docx)
            pdf_path  = storage.save_from_path(CATEGORY, pdf_name,  tmp_pdf)

        return {
            "docx_path": str(docx_path),
            "pdf_path":  str(pdf_path),
            "docx_filename": docx_name,
            "pdf_filename":  pdf_name,
        }

    def get_latest_paths(self, course_id) -> dict:
        storage = get_storage()
        docx = storage.get_latest(CATEGORY, prefix=f"qpaper_{course_id}_", ext=".docx")
        pdf  = storage.get_latest(CATEGORY, prefix=f"qpaper_{course_id}_", ext=".pdf")
        return {
            "docx_path": str(docx) if docx else None,
            "pdf_path":  str(pdf)  if pdf  else None,
        }

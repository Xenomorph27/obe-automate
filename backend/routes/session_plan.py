# backend/routes/session_plan.py
import json
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List

from fastapi import APIRouter, Depends, HTTPException
from backend.core.auth import require_auth
from backend.database.user_models import User
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession

from backend.core.exceptions import OBEException
from backend.core.logger import get_logger
from backend.core.storage import get_storage
from backend.database.connection import get_db
from backend.services.session_plan_service import SessionPlanService

logger = get_logger(__name__)
router = APIRouter(prefix="/session-plan", tags=["Session Plan"])

_CATEGORY = "session_plans"
_META_SUFFIX = "_edited.json"  # stores edited table state


class SavePlanRequest(BaseModel):
    cols: List[Dict[str, Any]]
    rows: List[Dict[str, Any]]


@router.post("/generate/{course_id}", status_code=201)
async def generate_session_plan(
    course_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """
    Generates a unit-wise session plan for the course using Gemini AI.
    Returns metadata and a download URL for the Word document.
    """
    logger.info(f"Session plan generation requested for course_id={course_id}")
    try:
        service = SessionPlanService(db)
        result = await service.generate(course_id)
        return {"status": "success", "data": result}
    except OBEException as e:
        logger.error(f"Session plan error: {e.message}")
        raise HTTPException(status_code=e.status_code, detail=e.message)
    except Exception:
        logger.exception(f"Unexpected error generating session plan for course {course_id}")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.get("/download/{course_id}")
async def download_session_plan(course_id: int, current_user: User = Depends(require_auth)):
    """Download the generated session plan Word document."""
    filepath = SessionPlanService.get_filepath(course_id)
    if not os.path.exists(filepath):
        raise HTTPException(
            status_code=404,
            detail=f"Session plan not found for course_id={course_id}. Run POST /session-plan/generate/{course_id} first.",
        )
    return FileResponse(
        path=filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"session_plan_{course_id}.docx",
    )


@router.post("/save/{course_id}", status_code=200)
async def save_session_plan(
    course_id: int,
    payload: SavePlanRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """
    Saves the user-edited table (cols + rows) for a session plan.
    Persists the data as JSON so it survives page reloads.
    Also rebuilds a fresh .docx from the edited rows.
    """
    storage = get_storage()

    # 1. Persist the edited table as JSON
    meta_filename = f"session_plan_{course_id}{_META_SUFFIX}"
    meta_bytes = json.dumps({"cols": payload.cols, "rows": payload.rows}, ensure_ascii=False).encode()
    storage.save(_CATEGORY, meta_filename, meta_bytes)

    # 2. Rebuild docx from the edited rows using python-docx
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        _NAVY = "1F3864"
        _WHITE = RGBColor(0xFF, 0xFF, 0xFF)

        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(0.6)
            sec.left_margin = sec.right_margin = Inches(0.7)

        # Header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Session Plan — Course ID {course_id} (Edited)")
        run.bold = True
        run.font.size = Pt(13)

        # Table
        col_labels = [c.get("label", c.get("key", "")) for c in payload.cols]
        tbl = doc.add_table(rows=1, cols=len(col_labels))
        tbl.style = "Table Grid"

        def _shade(cell, hex_color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            tcPr.append(shd)

        for i, (cell, label) in enumerate(zip(tbl.rows[0].cells, col_labels)):
            _shade(cell, _NAVY)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cell.paragraphs[0].add_run(label)
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = _WHITE

        _LIGHT = "D6DCE4"
        for ri, row_data in enumerate(payload.rows):
            row = tbl.add_row()
            if ri % 2 == 0:
                for c in row.cells:
                    _shade(c, _LIGHT)
            for ci, col in enumerate(payload.cols):
                val = str(row_data.get(col.get("key", ""), ""))
                cell = row.cells[ci]
                cell.paragraphs[0].clear()
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(8)

        docx_filename = f"session_plan_{course_id}.docx"
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp) / docx_filename
            doc.save(str(tmp_path))
            storage.save_from_path(_CATEGORY, docx_filename, tmp_path)

        logger.info(f"Session plan saved (edited) for course_id={course_id}: {len(payload.rows)} rows")
        return {
            "status": "success",
            "message": f"Saved {len(payload.rows)} rows",
            "download_url": f"/session-plan/download/{course_id}",
        }

    except Exception as exc:
        logger.exception(f"Failed to rebuild docx for course_id={course_id}: {exc}")
        # JSON was already saved — don't fail the whole request
        return {
            "status": "partial",
            "message": f"Data saved but docx rebuild failed: {exc}",
        }


@router.get("/materials/{course_id}")
async def get_session_materials(course_id: int, current_user: User = Depends(require_auth)):
    """
    Returns study materials extracted from the session plan.
    Reads the stored edited JSON first; falls back to the raw generated plan structure.
    """
    storage = get_storage()

    # Try edited JSON first
    meta_filename = f"session_plan_{course_id}{_META_SUFFIX}"
    meta_path = storage.get_path(_CATEGORY, meta_filename)

    textbooks: list = []
    web_links: list = []
    journals: list = []
    moocs: list = []

    if meta_path:
        try:
            data = json.loads(meta_path.read_text(encoding="utf-8"))
            rows = data.get("rows", [])
            cols = data.get("cols", [])
            # Look for material columns by label keywords
            for col in cols:
                label = col.get("label", "").lower()
                key = col.get("key", "")
                if any(k in label for k in ["textbook", "reference", "book"]):
                    textbooks += [{"title": r.get(key, ""), "author": "", "publisher": ""} for r in rows if r.get(key)]
                elif any(k in label for k in ["web", "link", "nptel", "url", "online"]):
                    web_links += [{"title": r.get(key, ""), "unit": r.get("unit", ""), "url": ""} for r in rows if r.get(key)]
                elif any(k in label for k in ["journal", "paper", "research"]):
                    journals += [{"title": r.get(key, ""), "url": ""} for r in rows if r.get(key)]
                elif any(k in label for k in ["mooc", "course", "swayam", "coursera"]):
                    moocs += [{"title": r.get(key, ""), "platform": "", "url": ""} for r in rows if r.get(key)]
        except Exception as e:
            logger.warning(f"Could not parse edited session plan JSON: {e}")

    # Provide curated defaults if nothing was found in table columns
    if not any([textbooks, web_links, journals, moocs]):
        raise HTTPException(
            status_code=404,
            detail="No study materials found. Generate the session plan first, then add material columns via AI chat (e.g. 'Add a Textbooks column').",
        )

    return {
        "status": "success",
        "data": {
            "textbooks": textbooks,
            "web_links": web_links,
            "journals": journals,
            "mooc_courses": moocs,
        },
    }
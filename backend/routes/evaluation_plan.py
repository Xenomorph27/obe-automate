# backend/routes/evaluation_plan.py
import json
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession

from backend.core.exceptions import OBEException
from backend.core.logger import get_logger
from backend.core.storage import get_storage
from backend.database.connection import get_db
from backend.services.evaluation_plan_service import EvaluationPlanService

logger = get_logger(__name__)
router = APIRouter(prefix="/evaluation-plan", tags=["Evaluation Plan"])

_CATEGORY = "evaluation_plans"
_META_SUFFIX = "_edited.json"


class SavePlanRequest(BaseModel):
    cols: List[Dict[str, Any]]
    rows: List[Dict[str, Any]]


@router.post("/generate/{course_id}", status_code=201)
async def generate_evaluation_plan(
    course_id: int,
    db: AsyncSession = Depends(get_db),
):
    """
    Generates a full CIE + SEE evaluation plan for the course using Gemini AI.
    Returns metadata and a download URL for the Word document.
    """
    logger.info(f"Evaluation plan generation requested for course_id={course_id}")
    try:
        service = EvaluationPlanService(db)
        result = await service.generate(course_id)
        return {"status": "success", "data": result}
    except OBEException as e:
        logger.error(f"Evaluation plan error: {e.message}")
        raise HTTPException(status_code=e.status_code, detail=e.message)
    except Exception:
        logger.exception(f"Unexpected error generating evaluation plan for course {course_id}")
        raise HTTPException(status_code=500, detail="Internal server error")


@router.get("/download/{course_id}")
async def download_evaluation_plan(course_id: int):
    """Download the generated evaluation plan Word document."""
    filepath = EvaluationPlanService.get_filepath(course_id)
    if not os.path.exists(filepath):
        raise HTTPException(
            status_code=404,
            detail=f"Evaluation plan not found for course_id={course_id}. Run POST /evaluation-plan/generate/{course_id} first.",
        )
    return FileResponse(
        path=filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"evaluation_plan_{course_id}.docx",
    )


@router.post("/save/{course_id}", status_code=200)
async def save_evaluation_plan(
    course_id: int,
    payload: SavePlanRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Saves the user-edited evaluation plan table (cols + rows).
    Persists as JSON and rebuilds a fresh .docx from the edited data.
    """
    storage = get_storage()

    # 1. Persist edited table as JSON
    meta_filename = f"evaluation_plan_{course_id}{_META_SUFFIX}"
    meta_bytes = json.dumps({"cols": payload.cols, "rows": payload.rows}, ensure_ascii=False).encode()
    storage.save(_CATEGORY, meta_filename, meta_bytes)

    # 2. Rebuild docx from edited rows
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        _NAVY = "1F3864"
        _WHITE = RGBColor(0xFF, 0xFF, 0xFF)
        _LIGHT = "D6DCE4"

        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(0.6)
            sec.left_margin = sec.right_margin = Inches(0.7)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Evaluation Plan — Course ID {course_id} (Edited)")
        run.bold = True
        run.font.size = Pt(13)

        col_labels = [c.get("label", c.get("key", "")) for c in payload.cols]
        tbl = doc.add_table(rows=1, cols=len(col_labels))
        tbl.style = "Table Grid"

        def _shade(cell, hex_color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            tcPr.append(shd)

        for cell, label in zip(tbl.rows[0].cells, col_labels):
            _shade(cell, _NAVY)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cell.paragraphs[0].add_run(label)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = _WHITE

        for ri, row_data in enumerate(payload.rows):
            row = tbl.add_row()
            if ri % 2 == 0:
                for c in row.cells:
                    _shade(c, _LIGHT)
            for ci, col in enumerate(payload.cols):
                val = str(row_data.get(col.get("key", ""), ""))
                cell = row.cells[ci]
                cell.paragraphs[0].clear()
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(9)

        docx_filename = f"evaluation_plan_{course_id}.docx"
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp) / docx_filename
            doc.save(str(tmp_path))
            storage.save_from_path(_CATEGORY, docx_filename, tmp_path)

        logger.info(f"Evaluation plan saved (edited) for course_id={course_id}: {len(payload.rows)} rows")
        return {
            "status": "success",
            "message": f"Saved {len(payload.rows)} rows",
            "download_url": f"/evaluation-plan/download/{course_id}",
        }

    except Exception as exc:
        logger.exception(f"Failed to rebuild eval plan docx for course_id={course_id}: {exc}")
        return {
            "status": "partial",
            "message": f"Data saved but docx rebuild failed: {exc}",
        }
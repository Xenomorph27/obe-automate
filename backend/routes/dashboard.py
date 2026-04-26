# backend/routes/dashboard.py
import io
import json
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from backend.core.logger import get_logger
from backend.core.storage import get_storage
from backend.database.connection import get_db
from backend.services.dashboard_service import DashboardService

logger = get_logger(__name__)
router = APIRouter(prefix="/dashboard", tags=["Dashboard"])


# ── Timetable Parser ────────────────────────────────────────────────────

def parse_timetable_docx(file_bytes: bytes) -> dict:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    result = {
        "faculty_name": "",
        "department": "",
        "academic_year": "",
        "time_slots": [],
        "schedule": {},
        "courses_taught": [],
    }
    if not doc.tables:
        raise ValueError("No table found in timetable document.")
    table = doc.tables[0]
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

    # Extract metadata
    for row in rows[:5]:
        text = row[0] if row else ""
        if "Department" in text or "AIML" in text:
            result["department"] = text
        elif "Individual Timetable" in text or "AY" in text:
            result["academic_year"] = text
        elif text and not any(k in text for k in ["Symbiosis", "Department", "Individual"]):
            if not result["faculty_name"]:
                result["faculty_name"] = text

    # Find header row
    header_idx = None
    time_slots = []
    for i, row in enumerate(rows):
        if row and row[0].strip().lower() in ("day/time", "day", "day\\time"):
            header_idx = i
            time_slots = [v for v in row[1:] if v.strip()]
            break
    if header_idx is None:
        for i, row in enumerate(rows):
            if any(":" in str(v) and "-" in str(v) for v in row):
                header_idx = i
                time_slots = [v for v in row[1:] if v.strip()]
                break

    result["time_slots"] = time_slots
    DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

    if header_idx is not None:
        current_day = None
        day_slots = {}
        for row in rows[header_idx + 1:]:
            if not row or not any(v for v in row):
                continue
            first_col = row[0].strip()
            matched_day = next((d for d in DAYS if first_col.lower().startswith(d.lower())), None)
            if matched_day:
                if current_day and day_slots:
                    result["schedule"][current_day] = _build_day_schedule(day_slots, time_slots)
                current_day = matched_day
                day_slots = {}
                for j, val in enumerate(row[1:len(time_slots)+1]):
                    if val.strip():
                        day_slots.setdefault(j, {})["course"] = val.strip()
            elif current_day:
                for j, val in enumerate(row[1:len(time_slots)+1]):
                    if val.strip():
                        existing = day_slots.setdefault(j, {})
                        if "course" not in existing:
                            existing["course"] = val.strip()
                        elif "section" not in existing:
                            existing["section"] = val.strip()
                        elif "room" not in existing:
                            existing["room"] = val.strip()
        if current_day and day_slots:
            result["schedule"][current_day] = _build_day_schedule(day_slots, time_slots)

    courses = []
    for row in rows:
        if len(row) >= 5 and any(kw in " ".join(row).lower() for kw in ["unsupervised","programming","learning","project","lab"]):
            for cell in row:
                if cell.strip() and len(cell.strip()) > 5 and cell.strip() not in courses:
                    if not any(t in cell for t in ["Professor","Faculty","SKO","FY","SY","TY"]):
                        courses.append(cell.strip())
    result["courses_taught"] = courses[:6]
    return result


def _build_day_schedule(day_slots: dict, time_slots: list) -> list:
    schedule = []
    for j, slot_time in enumerate(time_slots):
        if j in day_slots and day_slots[j].get("course"):
            schedule.append({
                "time": slot_time,
                "course": day_slots[j].get("course",""),
                "section": day_slots[j].get("section",""),
                "room": day_slots[j].get("room",""),
            })
    return schedule


# ── Routes ──────────────────────────────────────────────────────────────

@router.get("/department")
async def get_department_dashboard(db: AsyncSession = Depends(get_db)):
    svc = DashboardService(db)
    return await svc.get_department_summary()


@router.post("/timetable/upload", status_code=201)
async def upload_timetable(
    file: UploadFile = File(..., description="Individual timetable .docx file"),
):
    """Upload faculty individual timetable (.docx). Parses and stores it."""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    try:
        file_bytes = await file.read()
        timetable = parse_timetable_docx(file_bytes)
        storage = get_storage()
        storage.save("timetables", "current_timetable.json", json.dumps(timetable, ensure_ascii=False, indent=2).encode())
        logger.info(f"Timetable uploaded: {timetable.get('faculty_name','Unknown')}")
        return {"status": "success", "data": timetable}
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        logger.exception("Error parsing timetable")
        raise HTTPException(status_code=500, detail="Failed to parse timetable file.")


@router.get("/timetable")
async def get_timetable():
    """Get the currently uploaded timetable. Returns null if none uploaded."""
    try:
        storage = get_storage()
        path = storage.get_path("timetables", "current_timetable.json")
        if not path:
            return {"status": "success", "data": None}
        return {"status": "success", "data": json.loads(path.read_text(encoding="utf-8"))}
    except Exception:
        logger.exception("Error reading timetable")
        return {"status": "success", "data": None}